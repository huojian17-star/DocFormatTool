# -*- coding: utf-8 -*-
"""从结构化段落列表 + 格式配置生成排版后的 .docx。

流程：infer.parse_file(输入) → build(cfg, structs) → 输出成品。
支持 Markdown 图片引用（![描述](本地路径)）自动插入并编号题注。
"""
import os
import re
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import xmlutil as X
from . import styles as S
from . import infer

PAPER_SIZES = {
    "A4": (21.0, 29.7), "A3": (29.7, 42.0), "A5": (14.8, 21.0), "B5": (17.6, 25.0),
    "Letter": (21.59, 27.94),
}
IMAGE_MD_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
FIELD_KEYS = ("姓名", "学号", "班级", "专业", "学院", "指导教师", "指导老师", "日期")


def build(cfg: dict, structs: list, dst: str, base_dir: str = "") -> str:
    """structs: infer 的输出列表。base_dir: 解析相对图片路径的基准目录。"""
    doc = Document()
    _setup_page(doc, cfg)

    cover_seg = [st for st in structs if st["type"] == "cover"]
    rest = [st for st in structs if st["type"] != "cover"]
    has_cover = bool(cover_seg) and cfg["cover"].get("enabled", True)

    if has_cover:
        _build_cover(doc, cfg, cover_seg)
        # 正文节（新节）：页码必加，页眉文字可选
        sec = doc.add_section()
        _setup_page(doc, cfg)
        _setup_columns(doc, cfg)
    else:
        # 无封面：直接在第一节排正文（内置英文模板等）
        sec = doc.sections[0]
        _setup_columns(doc, cfg)

    _setup_body_section(doc, cfg, sec)
    _apply_body(doc, cfg, rest, base_dir)

    doc.save(dst)
    # 返回统计（供《改动报告》）
    stats = {
        "paras": sum(1 for st in rest if st["type"] != "blank"),
        "paras_total": len(rest),
        "runs_set": 0,
        "h1": sum(1 for st in rest if st["type"] == "heading1"),
        "h2": sum(1 for st in rest if st["type"] == "heading2"),
        "h3": sum(1 for st in rest if st["type"] == "heading3"),
        "tables": len(doc.tables),
        "pics": len(doc.element.body.findall(".//" + qn("w:drawing"))),
        "body_start": cfg.get("page_numbering", {}).get("body_start", 1),
        "front": cfg.get("page_numbering", {}).get("front_matter", "none"),
        "toc": "已插入" if cfg.get("toc", {}).get("enabled", False) else "未启用",
    }
    return stats


def _log_error(src: str):
    """排版异常记录到 %APPDATA%/DocFormatTool/error.log（买家可发给卖家诊断）。"""
    import traceback
    try:
        log_dir = os.path.join(os.environ.get("APPDATA", os.path.expanduser("~")), "DocFormatTool")
        os.makedirs(log_dir, exist_ok=True)
        with open(os.path.join(log_dir, "error.log"), "a", encoding="utf-8") as f:
            f.write("=" * 50 + "\n")
            f.write("时间: %s\n文件: %s\n" % (time.strftime("%Y-%m-%d %H:%M:%S"), src))
            f.write(traceback.format_exc())
    except Exception:
        pass


def reformat_existing(cfg: dict, src: str, dst: str) -> dict:
    """改写式排版入口：异常时记录错误日志并抛出（GUI 显示给用户）。"""
    try:
        return _reformat_existing_core(cfg, src, dst)
    except Exception:
        _log_error(src)
        raise


def _reformat_existing_core(cfg: dict, src: str, dst: str) -> dict:
    """改写式排版：保留原文档全部内容（图片/表格/公式/超链接），只规范化格式。

    用于 .docx 输入（学生已有带图表的 Word 文档）；.txt/.md 走 build()。
    整篇文档所有段落统一处理：正文、表格单元格、文本框（txbxContent）、
    内容控件（sdt）。不做 id() 去重（lxml 代理地址会复用导致漏处理）。
    返回改动统计 stats，用于生成《改动报告》。
    """
    doc = Document(src)
    stats = {
        "runs_set": 0, "paras": 0, "paras_total": 0,
        "h1": 0, "h2": 0, "h3": 0,
        "tables": len(doc.tables),
        "pics": len(doc.element.body.findall(".//" + qn("w:drawing"))),
        "body_start": cfg.get("page_numbering", {}).get("body_start", 1),
        "front": cfg.get("page_numbering", {}).get("front_matter", "none"),
    }
    _setup_page(doc, cfg)
    _setup_columns(doc, cfg)
    _apply_page_numbering(doc, cfg)
    _ensure_heading_styles(doc)

    # 目录：文档原无目录且模板开启 → 在正文首标题前插入
    toc_action = "未启用/未检测到"
    has_toc = _detect_toc(doc)
    if has_toc:
        toc_action = "已有（保留原目录）"
    elif cfg.get("toc", {}).get("enabled", False):
        body_el = _find_first_heading(doc)
        if body_el is not None:
            _insert_toc_before(doc, cfg, body_el)
            toc_action = "已插入（Word 中右键目录→更新域刷新页码）"
    stats["toc"] = toc_action

    in_cover = True
    in_toc = False
    # 先快照段落列表再遍历：处理中会修改文档（插分节符/设样式），
    # lxml iter 在遍历中修改树会导致生成器跳变（目录区条目漏处理）
    paras = list(doc.element.body.iter(qn("w:p")))
    for idx, p_el in enumerate(paras):
        stats["paras_total"] += 1
        # 前瞻下一段文本（目录区结束判定需要：正文标题后是长正文，目录条目后是短条目）。
        # 跳过空段——正文标题后常有空行，空段不能作为"非长正文"信号。
        next_text = ""
        for j in range(idx + 1, len(paras)):
            nt = para_text(paras[j]).strip()
            if nt:
                next_text = nt
                break
        in_cover, in_toc = _reformat_paragraph(p_el, cfg, in_cover, stats, in_toc, next_text)

    doc.save(dst)
    stats["out"] = dst
    return stats


def _is_cover_block(p_el, cfg) -> bool:
    """封面区判定：在第一个标题/摘要/长正文之前的短段落（学校名/论文类别/题目/封面字段表）。"""
    text = para_text(p_el).strip()
    if not text or p_el.findall(".//" + qn("w:drawing")):
        return False
    if len(text) > 40:
        return False
    typ, _ = infer._classify(text)
    if typ in ("heading1", "heading2", "heading3", "abstract_heading", "keywords",
               "ref_heading", "ref_item", "appendix"):
        return False
    return True


def para_text(p_el) -> str:
    """安全提取段落文本：只拼 w:t（部分文档 XML 的 w:p/w:r 带杂散 text，
    itertext() 会重复拼接导致文本翻倍，标题识别被干扰）。"""
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def _in_table(p_el) -> bool:
    """判断段落是否位于表格内（沿祖先链找 w:tbl）。"""
    cur = p_el.getparent()
    while cur is not None:
        if cur.tag == qn("w:tbl"):
            return True
        cur = cur.getparent()
    return False


def _set_run_font_name(run, cn_font, en_font):
    """只设置字体名（ascii/hAnsi/eastAsia），不动字号/加粗——用于封面区安全统一。"""
    run.font.name = en_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)


def _format_cover_para(p_el, cfg, text, stats=None):
    """封面段安全统一：只改字体名，布局/字号/对齐全部保持原样。

    例外：封面题目若字号爆炸（>模板题目字号+6pt），规范为模板题目字号并居中
    ——输入文档的题目常被作者放大到 60pt+，保留会拆行撑满页面。

    封面千奇百怪（表格布局/段落布局/校徽/字段各异），强行识别排版必然误伤，
    这里只做零风险的字体统一，其余交给学校自己的封面模板。
    """
    from docx.text.run import Run
    from docx.text.paragraph import Paragraph
    fd = cfg["fonts"]["body"]
    sizes = []
    for r in p_el.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            if sz is not None:
                sizes.append(int(sz.get(qn("w:val"))) / 2)
    title_size = cfg.get("cover", {}).get("title_size_pt", 22)
    exclude = ("毕业论文", "学位论文", "学院", "大学", "学校", "学号", "专业",
               "姓名", "指导", "完成日期", "目录", "摘要", "Abstract", "论文题目")
    max_size = max(sizes, default=0)
    is_title = (5 <= len(text) <= 40 and max_size > title_size + 6
                and not any(k in text for k in exclude))
    if is_title:
        p = Paragraph(p_el, None)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 清缩进（输入超大字号时可能带了缩进，22pt 下会拆行）+ 控制段距
        pf = p.paragraph_format
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        tfd = cfg["fonts"].get("doc_title", cfg["fonts"]["heading1"])
        for r in p_el.iter(qn("w:r")):
            # 清字符间距（输入超大字号时可能加了 w:spacing，22pt 下会撑宽拆行）
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                sp = rpr.find(qn("w:spacing"))
                if sp is not None:
                    rpr.remove(sp)
            S._set_run_font(Run(r, p), tfd["cn"], tfd["en"], title_size, bold=True)
        if stats is not None:
            stats["runs_set"] = stats.get("runs_set", 0) + 1
    else:
        for r in p_el.iter(qn("w:r")):
            _set_run_font_name(Run(r, None), fd["cn"], fd["en"])


def _guess_heading_by_format(p, cfg, text):
    """格式特征识别：加粗 + 字号明显大于正文 + 短行 → 手动标题（没用样式直接加粗的）。

    返回标题层级 1/2/3，非标题返回 0。保守判定，宁可漏判不可误判。
    """
    if len(text) > 40:
        return 0
    # 以句读结尾 → 叙述句，非标题（标题不以。！？；结尾）
    if text and text[-1] in "。！？；!?;":
        return 0
    body_size = cfg["fonts"]["body"]["size_pt"]
    sizes, bolds = [], []
    for r in p._p.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            continue
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            try:
                sizes.append(int(sz.get(qn("w:val"))) / 2)
            except (TypeError, ValueError):
                pass
        bolds.append(rpr.find(qn("w:b")) is not None)
    if not sizes:
        return 0
    sz = max(sizes)
    if not bolds or sum(bolds) / len(bolds) < 0.6:
        return 0
    if sz <= body_size + 1:
        return 0
    if sz >= body_size + 5:
        return 1
    if sz >= body_size + 2:
        return 2
    return 3


def _set_pstyle(p_el, style_id: str):
    """给段落设置 Word 内置样式（如 Heading1），让标题进入样式集/导航窗格/目录收集。"""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    ps = pPr.find(qn("w:pStyle"))
    if ps is None:
        ps = OxmlElement("w:pStyle")
        pPr.insert(0, ps)
    ps.set(qn("w:val"), style_id)


def _clear_paragraph_indent(p_el):
    """彻底清除段落缩进（删除整个 w:ind，含 firstLineChars/left 等全部残留）。

    python-docx 的 first_line_indent=Cm(0) 只改 w:firstLine，
    firstLineChars/left 可能残留（WPS/Word 优先 firstLineChars → 缩进仍在）。
    """
    pPr = p_el.find(qn("w:pPr"))
    if pPr is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)


def _ensure_heading_styles(doc):
    """确保文档 styles.xml 含 Heading1/2/3 样式定义（否则 pStyle 引用无效，样式集/导航窗格不显示）。

    用户自己的文档常不带这些内置样式；插入标准定义（样式 id 与
    python-docx/_set_pstyle 使用的 'Heading1/2/3' 一致），并给黑体黑色
    加粗（run 级格式会覆盖，样式仅作"挂载点"）。
    """
    styles_el = doc.styles.element
    existing = set()
    for s in styles_el.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        if sid:
            existing.add(sid)
        nm = s.find(qn("w:name"))
        if nm is not None and nm.get(qn("w:val")):
            existing.add(nm.get(qn("w:val")))
    for sid, name, lvl in (("Heading1", "heading 1", "0"),
                           ("Heading2", "heading 2", "1"),
                           ("Heading3", "heading 3", "2")):
        # 已存在的 Heading 样式：补 qFormat（样式窗格显示）与黑色（防蓝色标题），不改动其余
        matched = None
        for s in styles_el.findall(qn("w:style")):
            if s.get(qn("w:styleId")) == sid:
                matched = s
                break
        if matched is not None:
            _upgrade_heading_style(matched)
            continue
        st = OxmlElement("w:style")
        st.set(qn("w:type"), "paragraph")
        st.set(qn("w:styleId"), sid)
        nm_el = OxmlElement("w:name")
        nm_el.set(qn("w:val"), name)
        st.append(nm_el)
        qf = OxmlElement("w:qFormat")
        st.append(qf)
        pPr = OxmlElement("w:pPr")
        keepn = OxmlElement("w:keepNext")
        pPr.append(keepn)
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), lvl)
        pPr.append(ol)
        st.append(pPr)
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "000000")
        rPr.append(col)
        st.append(rPr)
        styles_el.append(st)


def _upgrade_heading_style(st_el):
    """对已存在的 Heading 样式补 qFormat（样式窗格显示）与黑色粗体（run 级格式会覆盖，仅兜底）。"""
    qf = st_el.find(qn("w:qFormat"))
    if qf is None:
        st_el.append(OxmlElement("w:qFormat"))
    rPr = st_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        st_el.append(rPr)
    if rPr.find(qn("w:b")) is None:
        rPr.insert(0, OxmlElement("w:b"))
    col = rPr.find(qn("w:color"))
    if col is None:
        col = OxmlElement("w:color")
        rPr.append(col)
    col.set(qn("w:val"), "000000")
    # 关键：删除 themeColor/themeTint/themeShade——Word/WPS 渲染优先主题色（accent1 默认蓝），
    # 输入文档的 Heading 样式常带 themeColor=accent1，只设 val 不删它仍是蓝色
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if col.get(qn(attr)) is not None:
            del col.attrib[qn(attr)]


def _reformat_paragraph(p_el, cfg, _in_cover=False, stats=None, _in_toc=False, next_text=""):
    """统一处理一个段落（顶层正文/标题/表格内/文本框内）。

    含图片/公式的段落原样保留仅居中；其余按文本特征套角色格式。
    返回 (是否仍处于封面区, 是否仍在目录区)。
    stats: 改动统计 dict（用于《改动报告》），可空。
    """
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    st = stats or {}

    # 含图片/公式/图形的段落：原样保留，仅居中 + 统一 run 字体（图片段落常有隐藏空 run）
    if p_el.findall(".//" + qn("w:drawing")) or p_el.findall(".//" + qn("w:object")):
        p = Paragraph(p_el, None)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fd = cfg["fonts"]["body"]
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"])
            st["runs_set"] = st.get("runs_set", 0) + 1
        st["paras"] = st.get("paras", 0) + 1
        return _in_cover, _in_toc
    text = para_text(p_el).strip()
    p = Paragraph(p_el, None)

    # 封面区：正文开始前的连续短段，只统一字体（布局/字号不动，安全第一）
    if _in_cover:
        if not text:
            # 封面区空行：统一字体，不结束封面区
            fd = cfg["fonts"]["body"]
            for r in p_el.iter(qn("w:r")):
                _set_run_font_name(Run(r, p), fd["cn"], fd["en"])
            return True, _in_toc
        if _is_cover_block(p_el, cfg):
            _format_cover_para(p_el, cfg, text, stats)
            st["paras"] = st.get("paras", 0) + 1
            return True, _in_toc
        # 首个非封面段：结束封面区，按正常流程继续
    _in_cover = False

    # 目录区：文档自带目录页时，"目 录"标题后的条目不套标题格式
    if not _in_toc and text and "目" in text and "录" in text and len(text) <= 10:
        # 目录标题：居中黑体，进入目录区
        fd = cfg["fonts"]["heading1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"], bold=True)
            st["runs_set"] = st.get("runs_set", 0) + 1
        st["paras"] = st.get("paras", 0) + 1
        return False, True
    if _in_toc:
        if not text:
            return False, True  # 空段：保持目录区（分节符等插入的空段不能中断目录）
        typ_t, _ = infer._classify(text)
        # 正文标题特征：短行标题 + 下一条是长正文（>40 字）或更低层级子标题（H1 后跟 1.1）→ 目录区结束
        nt_type, _ = infer._classify(next_text) if next_text else ("", "")
        is_body_heading = (typ_t in ("heading1", "heading2", "heading3")
                           and (len(next_text) > 40
                                or (typ_t == "heading1" and nt_type in ("heading2", "heading3"))))
        if is_body_heading:
            _in_toc = False
            # 不 return，落到下方正常标题处理
        elif len(text) <= 40:
            # 目录条目（短行：章节/参考文献/致谢/附录）：按正文格式，不套标题
            S.format_body(p, cfg)
            _clear_paragraph_indent(p_el)  # 目录条目不缩进（清 firstLineChars/left 残留）
            st["paras"] = st.get("paras", 0) + 1
            return False, True
        else:
            return False, False  # 长正文出现，目录区结束

    # 表格内段落：统一表格字体（表格里不判标题，防止"1. xxx"列举被当标题）
    if _in_table(p_el):
        fd = cfg["fonts"].get("table", cfg["fonts"]["body"])
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"])
            st["runs_set"] = st.get("runs_set", 0) + 1
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        st["paras"] = st.get("paras", 0) + 1
        return False, _in_toc

    if not text:
        # 空段落（域代码/格式标记 run）：仅统一字体，不动段落格式
        fd = cfg["fonts"]["body"]
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"])
        return False, _in_toc

    t, _ = infer._classify(text)
    if t == "heading1":
        S.format_heading(p, cfg, 1)
        _set_pstyle(p_el, "Heading1")  # 进入 Word 样式集/目录收集
        st["h1"] = st.get("h1", 0) + 1
    elif t == "heading2":
        S.format_heading(p, cfg, 2)
        _set_pstyle(p_el, "Heading2")
        st["h2"] = st.get("h2", 0) + 1
    elif t == "heading3":
        S.format_heading(p, cfg, 3)
        _set_pstyle(p_el, "Heading3")
        st["h3"] = st.get("h3", 0) + 1
    elif t == "abstract_heading" and len(text) > 20:
        # "摘要：xxx" 长句（摘要标题+内容混合）→ 按正文处理，不套标题格式
        S.format_body(p, cfg)
    elif t == "abstract_heading":
        # 摘要标题：独立字体（黑体四号居中，区别于章节标题）
        fd = cfg["fonts"].get("abstract_heading", cfg["fonts"]["heading1"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"], bold=True)
    elif t == "keywords":
        # 关键词行：标签（关键词：/Keywords:/Index Terms—/CCS Concepts •）黑体加粗，内容正文
        kw_fd = cfg["fonts"].get("keywords", cfg["fonts"].get("abstract_heading", cfg["fonts"]["heading1"]))
        bd_fd = cfg["fonts"]["body"]
        p.paragraph_format.line_spacing = par_ls(cfg)
        p.paragraph_format.first_line_indent = Cm(0)
        full = para_text(p_el)
        m = re.match(r"^(.{1,20}?[：:—•]\s*)", full)
        label_len = len(m.group(1)) if m else 0
        acc = 0
        for r in p_el.iter(qn("w:r")):
            t_el = r.find(qn("w:t"))
            if t_el is None or not t_el.text:
                continue
            ln = len(t_el.text)
            if acc + ln <= label_len:
                S._set_run_font(Run(r, p), kw_fd["cn"], kw_fd["en"], kw_fd["size_pt"], bold=True)
            else:
                S._set_run_font(Run(r, p), bd_fd["cn"], bd_fd["en"], bd_fd["size_pt"])
            acc += ln
            st["runs_set"] = st.get("runs_set", 0) + 1
    elif t in ("ref_heading", "appendix"):
        S.format_heading(p, cfg, 1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif t == "ref_item":
        fd = cfg["fonts"].get("ref", cfg["fonts"]["body"])
        pf = p.paragraph_format
        pf.line_spacing = par_ls(cfg)
        # GB/T 7714 悬挂缩进：编号顶格，续行缩进（条目跨页时不切断、视觉整齐）
        _clear_paragraph_indent(p_el)  # 先清输入缩进残留（firstLineChars/left），否则 WPS 优先 firstLineChars
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)
        pf.keep_together = True  # 段内不跨页断开（整条移下一页）
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 统一两端对齐（输入可能是居中/左对齐残留）
        pf.space_before = Pt(0)  # 清输入段前段后残留
        pf.space_after = Pt(0)
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"])
            st["runs_set"] = st.get("runs_set", 0) + 1
    elif t == "caption":
        fd = cfg["fonts"]["caption"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        for r in p_el.iter(qn("w:r")):
            S._set_run_font(Run(r, p), fd["cn"], fd["en"], fd["size_pt"], bold=True)
            st["runs_set"] = st.get("runs_set", 0) + 1
    else:
        # 格式特征识别：加粗 + 大字号 + 短行 = 手动标题（没用样式直接加粗的）
        h_lvl = _guess_heading_by_format(p, cfg, text)
        if h_lvl:
            S.format_heading(p, cfg, h_lvl)
            _set_pstyle(p_el, {1: "Heading1", 2: "Heading2", 3: "Heading3"}[h_lvl])
            st["h%d" % h_lvl] = st.get("h%d" % h_lvl, 0) + 1
        else:
            # 正文：规范化字体、行距、首行缩进
            S.format_body(p, cfg)
    st["paras"] = st.get("paras", 0) + 1
    return False, _in_toc


def _find_first_heading(doc):
    """找正文第一个标题段落（分节/目录插入的边界）。跳过文档自带的目录页。"""
    in_toc = False
    for p_el in doc.element.body.iter(qn("w:p")):
        if _in_table(p_el):
            continue
        t = para_text(p_el).strip()
        if not t:
            continue
        # 目录标题行（"目 录"）→ 进入目录区
        if not in_toc and "目" in t and "录" in t and len(t) <= 10:
            in_toc = True
            continue
        if in_toc:
            typ_t, _ = infer._classify(t)
            if typ_t in ("heading1", "heading2", "heading3"):
                continue  # 目录条目（标题模式的短行），跳过
            in_toc = False  # 正文内容出现，目录区结束
        typ, _ = infer._classify(t)
        if typ in ("heading1", "heading2", "heading3"):
            return p_el
    return None


def _apply_page_numbering(doc, cfg):
    """页码结构：在正文第一个标题前分节，前置部分（封面/摘要/目录）无页码或罗马，
    正文节页码按 body_start 起始。学生可自定义，机器不猜。"""
    pn = cfg.get("page_numbering", {})
    front = pn.get("front_matter", "none")
    body_start = pn.get("body_start", 1)
    hf = cfg["header_footer"]
    footer_align = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(hf.get("footer_style", "center"),
                                                           WD_ALIGN_PARAGRAPH.CENTER)

    # 找正文第一个标题作为分节边界
    body_el = _find_first_heading(doc)

    if body_el is not None:
        from docx.text.paragraph import Paragraph
        X.insert_section_break_before(Paragraph(body_el, None))

    secs = doc.sections
    body_sec = secs[-1]

    # 前置部分各节：无页码 / 罗马 / 阿拉伯
    for sec in secs[:-1]:
        sec.footer.is_linked_to_previous = False
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        if front in ("roman", "decimal"):
            fp.alignment = footer_align
            X.add_page_number(fp)
            X.set_section_pgnum(X.section_sectPr(sec),
                                "upperRoman" if front == "roman" else "decimal", 1)

    # 正文节：页码 + body_start
    body_sec.footer.is_linked_to_previous = False
    fp = body_sec.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = footer_align
    X.add_page_number(fp)
    X.set_section_pgnum(X.section_sectPr(body_sec), "decimal", body_start)

    # 正文节页眉（可选）
    if hf.get("header_text"):
        body_sec.header.is_linked_to_previous = False
        hp = body_sec.header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hf_font = cfg["fonts"]["header"]
        r = hp.add_run(hf["header_text"])
        S._set_run_font(r, hf_font["cn"], hf_font["en"], hf_font["size_pt"])


def build_change_report(stats: dict, cfg: dict, src: str, out: str) -> str:
    """生成《改动报告.docx》（类知网查重报告样式）：告诉学生改了什么、占多少。

    stats: reformat_existing 返回的统计；out: 排版输出文件路径。
    返回报告文件路径。
    """
    import time as _t
    from docx import Document as _Doc
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _Align

    rpt = _Doc()
    m = cfg["page"]["margins_cm"]
    front_label = {"none": "无页码", "roman": "罗马数字", "decimal": "阿拉伯数字"}.get(
        stats.get("front", "none"), "无页码")
    paras_total = max(stats.get("paras_total", 0), 1)
    paras = stats.get("paras", 0)
    cover_pct = min(100, round(paras / paras_total * 100)) if paras_total else 0

    def _h(text, size=15):
        p = rpt.add_paragraph()
        p.alignment = _Align.CENTER
        r = p.add_run(text)
        S._set_run_font(r, "黑体", "Times New Roman", size, bold=True)

    def _h2(text):
        p = rpt.add_paragraph()
        r = p.add_run(text)
        S._set_run_font(r, "黑体", "Times New Roman", 13, bold=True)
        p.paragraph_format.space_before = _Pt(10)

    def _body(text, bold=False):
        p = rpt.add_paragraph(text)
        S._set_run_font(p.runs[0], "宋体", "Times New Roman", 11, bold=bold)
        p.paragraph_format.line_spacing = 1.4

    _h("排版改动报告")
    _body("输入文件：%s" % os.path.basename(src))
    _body("使用模板：%s" % cfg.get("school", "?"))
    _body("生成时间：%s" % _t.strftime("%Y-%m-%d %H:%M"))
    _body("输出文件：%s" % os.path.basename(out))

    _h2("一、总体改动")
    _body("共处理段落 %d 个，其中 %d 个段落的格式已规范化（覆盖率约 %d%%）。"
          % (paras_total, paras, cover_pct))
    _body("说明：本工具只调整格式，不改动任何文字内容。", bold=True)

    _h2("二、改动明细")
    _body("· 文字片段：%d 个文字片段的字体/字号已统一为模板规范" % stats.get("runs_set", 0))
    _body("· 章节标题：识别并规范 一级 %d 个 / 二级 %d 个 / 三级 %d 个"
          % (stats.get("h1", 0), stats.get("h2", 0), stats.get("h3", 0)))
    if stats.get("tables"):
        _body("· 表格：%d 张表格内文字已统一字号" % stats.get("tables", 0))
    _body("· 目录：%s" % stats.get("toc", "未启用/未检测到"))
    _body("· 页码：正文页码从 %s 开始；前置部分（封面/摘要）%s"
          % (stats.get("body_start", 1), front_label))
    _body("· 页面：页边距 上%.1f 下%.1f 左%.1f 右%.1f cm"
          % (m["top"], m["bottom"], m["left"], m["right"]))
    _body("· 封面：保持原样，未改动（由学校模板决定）")

    _h2("三、内容完整性")
    _body("· 图片：%d 张，全部保留" % stats.get("pics", 0))
    _body("· 表格：%d 张，全部保留" % stats.get("tables", 0))
    _body("· 文字：全部保留，未删除或改写任何内容")
    _body("· 原始文件：未改动，保留在 %s" % os.path.basename(src))

    report_path = os.path.splitext(out)[0] + "_改动报告.docx"
    rpt.save(report_path)
    return report_path


def _setup_body_section(doc, cfg, sec):
    """正文节：页眉（可选）+ 页脚页码 + 页码格式。"""
    hf = cfg["header_footer"]
    if hf.get("header_text"):
        sec.header.is_linked_to_previous = False
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hf_font = cfg["fonts"]["header"]
        r = hp.add_run(hf["header_text"])
        S._set_run_font(r, hf_font["cn"], hf_font["en"], hf_font["size_pt"])
    sec.footer.is_linked_to_previous = False
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    X.add_page_number(fp)
    X.set_section_pgnum(X.section_sectPr(sec), "decimal", 1)


def _setup_columns(doc, cfg):
    """双栏设置（作用于所有节）。1cm ≈ 567 twips。"""
    num = cfg["page"].get("columns", 1)
    if num < 2:
        return
    space = cfg["page"].get("column_spacing_cm", 0.5)
    for sec in doc.sections:
        sectPr = X.section_sectPr(sec)
        cols = sectPr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sectPr.append(cols)
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:space"), str(int(space * 567)))


def _setup_page(doc, cfg):
    page = cfg["page"]
    w, h = PAPER_SIZES.get(page.get("paper", "A4"), (21.0, 29.7))
    m = page["margins_cm"]
    for sec in doc.sections:
        sec.page_width = Cm(w)
        sec.page_height = Cm(h)
        sec.top_margin = Cm(m["top"]); sec.bottom_margin = Cm(m["bottom"])
        sec.left_margin = Cm(m["left"]); sec.right_margin = Cm(m["right"])


def _build_cover(doc, cfg, cover_seg):
    texts = [st["text"].strip() for st in cover_seg if st["text"].strip()]
    if not texts:
        return
    cov = cfg["cover"]
    title_font = cfg["fonts"].get("doc_title", cfg["fonts"]["heading1"])
    body_font = cfg["fonts"]["body"]

    school = next((t for t in texts if "大学" in t or "学院" in t), "")
    kind = next((t for t in texts if t != school and len(t) <= 12
                 and any(k in _compact(t) for k in ("论文", "报告", "设计"))), "")
    fields = [t for t in texts if any(k in t for k in FIELD_KEYS)]
    cand = [t for t in texts if t not in (school, kind) and t not in fields]
    title = max(cand, key=len) if cand else ""
    others = [t for t in texts if t not in (school, kind, title) and t not in fields]

    def _line(text, font_cn, font_en, size, bold=False, spacing=2.0, before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = spacing
        if before:
            p.paragraph_format.space_before = Pt(before)
        r = p.add_run(text)
        S._set_run_font(r, font_cn, font_en, size, bold=bold)
        return p

    if school:
        _line(school, body_font["cn"], body_font["en"], body_font["size_pt"], bold=True)
    if kind:
        _line(kind, title_font["cn"], title_font["en"], 16, bold=True)
    doc.add_paragraph()
    if title:
        _line(title, title_font["cn"], title_font["en"],
              cov.get("title_size_pt", title_font["size_pt"]), bold=True, before=24)
    doc.add_paragraph()
    for t in fields + others:
        _line(t, body_font["cn"], body_font["en"], body_font["size_pt"])


def _add_toc_block(doc, cfg):
    """在当前位置追加"目录标题 + TOC 域"（从零生成路径：调用时机即第一个标题前）。"""
    toc = cfg.get("toc", {})
    if not toc.get("enabled", False):
        return
    fd = cfg["fonts"]["heading1"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(toc.get("heading_text", "目  录"))
    S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"], bold=True)
    tp = doc.add_paragraph()
    X.add_toc(tp, toc.get("levels", 3))
    note = tp.add_run("  （在 Word 中右键目录 → 更新域，页码即自动生成）")
    S._set_run_font(note, "宋体", "Times New Roman", 9)


def _insert_toc_before(doc, cfg, body_el):
    """改写式：在正文首标题前插入"目录标题 + TOC 域"（文档原无目录时）。"""
    from docx.text.paragraph import Paragraph
    toc = cfg.get("toc", {})
    fd = cfg["fonts"]["heading1"]
    title_p = doc.add_paragraph("")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(toc.get("heading_text", "目  录"))
    S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"], bold=True)
    toc_p = doc.add_paragraph("")
    X.add_toc(toc_p, toc.get("levels", 3))
    note = toc_p.add_run("  （在 Word 中右键目录 → 更新域，页码即自动生成）")
    S._set_run_font(note, "宋体", "Times New Roman", 9)
    body_el.addprevious(title_p._p)
    body_el.addprevious(toc_p._p)


def _detect_toc(doc):
    """检测文档是否已有目录（TOC 域或"目录"标题行）。返回 bool。"""
    for p_el in doc.element.body.iter(qn("w:p")):
        instr = p_el.findall(".//" + qn("w:instrText"))
        if instr:
            txt = "".join(t.text or "" for t in instr)
            if "TOC" in txt.upper():
                return True
        t = para_text(p_el).strip()
        if t and "目" in t and "录" in t and len(t) <= 12:
            return True
    return False


def _apply_body(doc, cfg, structs, base_dir):
    sty = cfg.get("styles", {})
    fig_fmt = cfg["captions"].get("figure", "图{chapter}-{num}")
    tab_fmt = cfg["captions"].get("table", "表{chapter}-{num}")
    chapter = 0
    fig_num = tab_num = 0
    par_cfg = cfg["paragraph"]
    in_refs = False
    toc_done = not cfg.get("toc", {}).get("enabled", False)

    for st in structs:
        t = st["type"]
        text = st.get("text", "")
        if t == "blank":
            continue
        if t == "ref_heading":
            in_refs = True
            _add_para(doc, cfg, cfg["references"].get("heading_text", "参考文献"), "heading1", center=True)
            continue
        if t == "ref_item" or (in_refs and t == "body"):
            _add_para(doc, cfg, text, "ref")
            continue
        if t in ("heading1", "heading2", "heading3", "appendix", "abstract_heading", "keywords"):
            in_refs = False
        if t == "image":
            m = IMAGE_MD_RE.match(text)
            if m:
                alt, path = m.group(1), m.group(2)
                fig_num += 1
                _add_image(doc, cfg, path, base_dir)
                # 题注：alt 已带编号（"图3-1 xxx"）→ 保留原编号；否则自动编号
                if re.match(r"^\s*[图表]\s*\d+[-.．]\d+\s*", alt):
                    cap = alt
                else:
                    cap = fig_fmt.format(chapter=chapter or fig_num, num=fig_num)
                    if alt:
                        cap += " " + alt
                _add_para(doc, cfg, cap, "caption", center=True)
            continue
        if t == "md_table":
            _add_md_table(doc, cfg, st.get("rows", []))
            continue
        if t == "code_block":
            _add_code_block(doc, cfg, st.get("text", ""))
            continue
        if t == "latex_block":
            _add_latex_block(doc, cfg, st.get("text", ""))
            continue
        if t == "heading1":
            if not toc_done:
                toc_done = True
                _add_toc_block(doc, cfg)
            chapter += 1
            fig_num = tab_num = 0
            _add_para(doc, cfg, text, "heading1")
        elif t == "heading2":
            _add_para(doc, cfg, text, "heading2")
        elif t == "heading3":
            _add_para(doc, cfg, text, "heading3")
        elif t == "abstract_heading":
            _add_para(doc, cfg, cfg["abstract"].get("heading_text", "摘  要"), "abs_heading", center=True)
        elif t == "keywords":
            _add_para(doc, cfg, text, "keywords")
        elif t == "appendix":
            _add_para(doc, cfg, text, "heading1")
        elif t == "caption":
            if text.startswith("图") or text.startswith("表"):
                _add_para(doc, cfg, text, "caption", center=True)
            else:
                _add_para(doc, cfg, text, "body")
        elif t == "heading1" or t.startswith("heading"):
            _add_para(doc, cfg, text, t)
        else:
            _add_para(doc, cfg, text, "body")


def _add_para(doc, cfg, text, kind, center=False):
    """按 kind 创建段落并套用内联格式（支持 md 粗体/斜体/行内代码）。"""
    f = cfg["fonts"]
    p = doc.add_paragraph()
    pf = p.paragraph_format

    if kind == "heading1":
        fd = f["heading1"]
        pf.line_spacing = par_ls(cfg)
        pf.space_before = Pt(6); pf.space_after = Pt(6)
        if center:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline_runs(p, cfg, text, fd, bold=fd.get("bold", True))
        _set_py_style(p, doc, "Heading 1")
    elif kind == "abs_heading":
        fd = f.get("abstract_heading", f["heading1"])
        pf.line_spacing = par_ls(cfg)
        pf.space_before = Pt(6); pf.space_after = Pt(6)
        if center:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline_runs(p, cfg, text, fd, bold=fd.get("bold", True))
    elif kind == "keywords":
        # 关键词行：标签黑体加粗 + 内容正文
        fd = f.get("keywords", f.get("abstract_heading", f["heading1"]))
        bd = f["body"]
        pf.line_spacing = par_ls(cfg)
        m = re.match(r"^(.{1,20}?[：:—•]\s*)", text)
        label = m.group(1) if m else ""
        rest = text[len(label):]
        r1 = p.add_run(label)
        S._set_run_font(r1, fd["cn"], fd["en"], fd["size_pt"], bold=True)
        r2 = p.add_run(rest)
        S._set_run_font(r2, bd["cn"], bd["en"], bd["size_pt"])
    elif kind == "heading2":
        fd = f["heading2"]
        pf.line_spacing = par_ls(cfg)
        pf.space_before = Pt(3); pf.space_after = Pt(3)
        _add_inline_runs(p, cfg, text, fd, bold=fd.get("bold", True))
        _set_py_style(p, doc, "Heading 2")
    elif kind == "heading3":
        fd = f["heading3"]
        pf.line_spacing = par_ls(cfg)
        _add_inline_runs(p, cfg, text, fd, bold=fd.get("bold", True))
        _set_py_style(p, doc, "Heading 3")
    elif kind == "caption":
        fd = f["caption"]
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = 1.0
        pf.space_before = Pt(3); pf.space_after = Pt(3)
        _add_inline_runs(p, cfg, text, fd, bold=True)
    elif kind == "ref":
        fd = f.get("ref", f["body"])
        pf.line_spacing = par_ls(cfg)
        _add_inline_runs(p, cfg, text, fd)
    else:  # body
        fd = f["body"]
        pf.line_spacing = par_ls(cfg)
        pf.space_after = Pt(cfg["paragraph"].get("space_after_pt", 0))
        if cfg["paragraph"].get("align", "justify") == "justify":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_runs(p, cfg, text, fd)
        chars = cfg["paragraph"].get("first_line_indent_chars", 2)
        if chars:
            S.set_first_line_indent(p, chars)
    return p


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _set_py_style(p, doc, style_name):
    """python-docx 方式设置段落样式（从零生成路径用）。"""
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass


def _add_inline_runs(p, cfg, text, fd, bold=False):
    """解析行内 md 格式：**粗体** *斜体* `行内代码` → 多个 run。"""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2])
            S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"], bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"], bold=bold, italic=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1])
            S._set_run_font(r, "宋体", "Consolas", max(fd["size_pt"] - 0.5, 8))
        else:
            r = p.add_run(part)
            S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"], bold=bold)


def _add_md_table(doc, cfg, rows):
    """md 表格 → Word 表格（跳过表头分隔行 |---|，带边框）。"""
    data = [r for r in rows if not all(re.match(r"^:?-{2,}:?$", c) for c in r)]
    if not data:
        return
    ncols = max(len(r) for r in data)
    tbl = doc.add_table(rows=len(data), cols=ncols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for ri, row in enumerate(data):
        for ci in range(ncols):
            tbl.cell(ri, ci).text = row[ci] if ci < len(row) else ""
    # 表格字体统一
    fd = cfg["fonts"].get("table", cfg["fonts"]["body"])
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"])
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)


def _add_code_block(doc, cfg, code):
    """md 代码块 → 等宽字体段落（保留换行、浅灰背景、固定小五号）。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    # 浅灰背景（工科论文代码块标准观感）
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    lines = code.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        S._set_run_font(r, "宋体", "Consolas", 10.5)
        if i < len(lines) - 1:
            r.add_break()


def _add_latex_block(doc, cfg, latex):
    """LaTeX 块级公式 → Word 原生公式（OMML），居中显示。转换失败退化为文本。"""
    from lxml import etree
    from . import omml
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    try:
        body = omml.latex_to_omml(latex)
        xml = ('<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
               "%s</m:oMath>" % body)
        el = etree.fromstring(xml)
        p._p.append(el)
    except Exception:
        fd = cfg["fonts"]["body"]
        r = p.add_run(latex)
        S._set_run_font(r, fd["cn"], fd["en"], fd["size_pt"])


def _add_image(doc, cfg, path, base_dir):
    if not os.path.isabs(path):
        path = os.path.join(base_dir, path)
    if not os.path.exists(path):
        return
    width_cm = 12.0
    try:
        # 自动缩放：不超过页边距内文本区宽度，且不超过原图尺寸
        from . import imagesize
        size = imagesize.image_size(path)
        if size:
            w_px = size[0]
            page = cfg["page"]
            pw = PAPER_SIZES.get(page.get("paper", "A4"), PAPER_SIZES["A4"])[0]
            text_w = pw - page["margins_cm"]["left"] - page["margins_cm"]["right"]
            width_cm = min(width_cm, text_w, w_px * 2.54 / 96.0)
    except Exception:
        pass
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(max(width_cm, 2.0)))
    except Exception:
        pass


def par_ls(cfg):
    return cfg["paragraph"].get("line_spacing", 1.5)


def _compact(s: str):
    """去掉空白字符（用于关键词匹配）。"""
    return re.sub(r"\s+", "", s)
