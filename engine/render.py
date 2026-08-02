# -*- coding: utf-8 -*-
"""排版主流程：读草稿模板 → 套学校格式 → 插入封面/目录/页眉页脚 → 另存成品。

不改内容结构（图片、表格、文字原样保留），只动格式与版式元素。
"""
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import styles as S
from . import xmlutil as X

PAPER_SIZES = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "B5": (17.6, 25.0),
}

_CN_NUM = "一二三四五六七八九十"


def _cn_number(n: int) -> str:
    if n <= 10:
        return _CN_NUM[n - 1]
    return str(n)


def render(cfg: dict, src: str, dst: str) -> str:
    """src: 学生填好的草稿模板.docx；dst: 输出成品路径。返回 dst。"""
    doc = Document(src)
    sty = cfg["styles"]
    paras = list(doc.paragraphs)

    # ---------- 1. 页面设置（所有节统一纸张与边距） ----------
    page = cfg["page"]
    w_cm, h_cm = PAPER_SIZES.get(page.get("paper", "A4"), PAPER_SIZES["A4"])
    m = page["margins_cm"]
    for sec in doc.sections:
        sec.page_width = Cm(w_cm)
        sec.page_height = Cm(h_cm)
        sec.top_margin = Cm(m["top"])
        sec.bottom_margin = Cm(m["bottom"])
        sec.left_margin = Cm(m["left"])
        sec.right_margin = Cm(m["right"])

    # ---------- 2. 删除"使用说明"段落 ----------
    for p in list(paras):
        if p.style.name == sty.get("instructions"):
            p._p.getparent().remove(p._p)

    # ---------- 3. 分节：封面 | 摘要+目录 | 正文 ----------
    abs_head = _first_para(paras, sty.get("abstract_heading"))
    first_h1 = _first_para(paras, sty.get("heading1"))
    if abs_head is not None:
        X.insert_section_break_before(abs_head)
    if first_h1 is not None:
        X.insert_section_break_before(first_h1)

    # ---------- 4. 页眉页脚与页码 ----------
    _setup_header_footer(doc, cfg)

    # ---------- 5. 封面格式化 ----------
    _format_cover(doc, cfg)

    # ---------- 6. 目录标题 + TOC 域 ----------
    _setup_toc(doc, cfg)

    # ---------- 7. 摘要标题 / 关键词 ----------
    _format_abstract(doc, cfg)

    # ---------- 8. 正文各样式套格式 + 题注编号 ----------
    _format_body_and_captions(doc, cfg)

    doc.save(dst)
    return dst


# ---------------------------------------------------------------- helpers

def _first_para(paras, style_name):
    if not style_name:
        return None
    for p in paras:
        if p.style.name == style_name:
            return p
    return None


def _setup_header_footer(doc, cfg):
    hf = cfg["header_footer"]
    font = cfg["fonts"]["header"]
    header_text = hf.get("header_text", "")
    secs = doc.sections

    def _put_header(section):
        if not header_text:
            return
        section.header.is_linked_to_previous = False
        par = section.header.paragraphs[0]
        par.text = ""
        align = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                 "left": WD_ALIGN_PARAGRAPH.LEFT,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT}[hf.get("header_align", "center")]
        par.alignment = align
        run = par.add_run(header_text)
        S._set_run_font(run, font["cn"], font["en"], font["size_pt"])

    def _put_page_footer(section, fmt, start):
        section.footer.is_linked_to_previous = False
        par = section.footer.paragraphs[0]
        par.text = ""
        align = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT}[hf.get("footer_style", "center")]
        par.alignment = align
        X.add_page_number(par)
        X.set_section_pgnum(X.section_sectPr(section), fmt, start)

    if len(secs) >= 3:
        # 封面节：保持空页眉页脚
        secs[1].header.is_linked_to_previous = False
        secs[1].footer.is_linked_to_previous = False
        _put_header(secs[1])
        _put_header(secs[2])
        _put_page_footer(secs[1], "upperRoman", 1)
        _put_page_footer(secs[2], "decimal", 1)
    else:
        # 兜底：单节场景只加正文页脚
        _put_header(secs[0])
        _put_page_footer(secs[0], "decimal", 1)


def _format_cover(doc, cfg):
    cov = cfg["cover"]
    sty = cfg["styles"]
    if not cov.get("enabled", True):
        return
    title_font = cfg["fonts"]["heading1"]
    body_font = cfg["fonts"]["body"]
    for p in doc.paragraphs:
        name = p.style.name
        p.paragraph_format.line_spacing = cov.get("line_spacing", 1.5)
        p.paragraph_format.space_after = Pt(6)
        if name == sty.get("cover_title"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                S._set_run_font(run, title_font["cn"], title_font["en"],
                                cov.get("title_size_pt", 26), bold=True)
        elif name == sty.get("cover_field"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                S._set_run_font(run, body_font["cn"], body_font["en"], body_font["size_pt"])


def _setup_toc(doc, cfg):
    sty = cfg["styles"]
    toc = cfg["toc"]
    heading_font = cfg["fonts"]["heading1"]
    for p in doc.paragraphs:
        if p.style.name == sty.get("toc_heading"):
            p.text = toc.get("heading_text", "目  录")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                S._set_run_font(run, heading_font["cn"], heading_font["en"],
                                heading_font["size_pt"], bold=True)
        elif p.style.name == sty.get("toc_placeholder"):
            p.text = ""
            S.set_paragraph_format(p, cfg["paragraph"])
            X.add_toc(p, toc.get("levels", 3))
            if toc.get("need_refresh_note", True):
                note = p.add_run("  （在 Word 中右键目录 → 更新域，页码即自动生成）")
                S._set_run_font(note, "宋体", "Times New Roman", 9)


def _format_abstract(doc, cfg):
    sty = cfg["styles"]
    abs_cfg = cfg["abstract"]
    heading_font = cfg["fonts"]["heading1"]
    body_font = cfg["fonts"]["body"]
    for p in doc.paragraphs:
        name = p.style.name
        if name == sty.get("abstract_heading"):
            p.text = abs_cfg.get("heading_text", "摘  要")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                S._set_run_font(run, heading_font["cn"], heading_font["en"],
                                heading_font["size_pt"], bold=True)
        elif name == sty.get("keywords"):
            label = abs_cfg.get("keywords_label", "关键词：")
            text = p.text or ""
            if not text.startswith(label):
                p.text = label + text
            S.format_body(p, cfg)
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                S._set_run_font(run, body_font["cn"], body_font["en"], body_font["size_pt"])


def _format_body_and_captions(doc, cfg):
    sty = cfg["styles"]
    fig_fmt = cfg["captions"].get("figure", "图{chapter}-{num}")
    tab_fmt = cfg["captions"].get("table", "表{chapter}-{num}")

    chapter = 0
    fig_num = 0
    tab_num = 0
    fig_seen = tab_seen = False

    for p in doc.paragraphs:
        name = p.style.name
        if name == sty.get("heading1"):
            chapter += 1
            fig_num = tab_num = 0
            S.format_heading(p, cfg, 1)
        elif name == sty.get("heading2"):
            S.format_heading(p, cfg, 2)
        elif name == sty.get("heading3"):
            S.format_heading(p, cfg, 3)
        elif name == sty.get("figure_caption"):
            fig_num += 1
            fig_seen = True
            desc = _caption_desc(p.text or "", "图")
            num = _caption_number(chapter, fig_num)
            new_text = fig_fmt.format(chapter=chapter or num, num=fig_num)
            S.format_caption(p, cfg, new_text + (" " + desc if desc else ""))
        elif name == sty.get("table_caption"):
            tab_num += 1
            tab_seen = True
            desc = _caption_desc(p.text or "", "表")
            new_text = tab_fmt.format(chapter=chapter or tab_num, num=tab_num)
            S.format_caption(p, cfg, new_text + (" " + desc if desc else ""))
        elif name == sty.get("ref_heading"):
            refs = cfg["references"]
            p.text = refs.get("heading_text", "参考文献")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                f = cfg["fonts"]["heading1"]
                S._set_run_font(run, f["cn"], f["en"], f["size_pt"], bold=True)
        elif name == sty.get("ref_item"):
            S.format_ref_item(p, cfg)
        elif name == sty.get("body") or name == "Normal" or name == "正文":
            S.format_body(p, cfg)
        elif name == sty.get("abstract_body"):
            S.format_body(p, cfg)


def _caption_number(chapter, num):
    """题注编号：有章号用"章-序号"，无章号退化为顺序号。"""
    if chapter:
        return "%d-%d" % (chapter, num)
    return str(num)


def _caption_desc(text, prefix):
    """去掉题注段落的"图/表"前缀与已有编号（如"图3-1"），返回纯描述文字。"""
    t = (text or "").strip()
    t = t.lstrip(prefix).lstrip(" ")
    t = re.sub(r"^\d+[-\.]\d+\s*", "", t)
    return t.strip()
