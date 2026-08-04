# -*- coding: utf-8 -*-
"""LoadLibrary 问题描述生成 docx（含用户新发现的 DocFormatTool_new 残留线索）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 全局字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')

# 标题
h = doc.add_heading('', level=0)
run = h.add_run('PyInstaller onefile 更新后 LoadLibrary 失败 —— 问题描述（求论坛大神诊断）')
run.font.name = '微软雅黑'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)

def para(text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = '微软雅黑'
    if color:
        r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = '微软雅黑'
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    return p

def h1(text):
    p = doc.add_heading('', level=1)
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x00, 0x50, 0x00)
    return p

# ============ 环境 ============
h1('一、环境')
for t in [
    'Windows 11 家庭版 中文版 10.0.26200，AMD64',
    'Python 3.13.9（conda 环境 D:\\Aconnada）',
    'PyInstaller 6.21.0',
    '打包方式：onefile，console=False，upx=True，自定义 spec（含 configs 目录 datas + icon）',
    '无任何杀软（防火墙已关闭，Windows Defender 相关也关闭）',
    '程序：tkinter 桌面工具（DocFormatTool），自带单实例锁（CreateMutexW）',
    '更新机制：程序内点"自动更新" → 后台线程下载新版 exe 到同目录 DocFormatTool_new.exe → 写 update.bat → os.startfile(bat) → 主程序 self.destroy 退出 → bat 执行替换 → 启动新版',
]:
    bullet(t)

# ============ 现象 ============
h1('二、现象（必现）')
for t in [
    '1. 双击运行 v1.0.12（onefile exe）→ 正常打开，能点"自动更新"',
    '2. 点自动更新 → 下载 v1.0.13 到 DocFormatTool_new.exe（下载文件 md5 已验证与发布源一致，没有损坏）',
    '3. update.bat 执行：taskkill 杀主程序 → copy 覆盖 → 延迟 → start "" "exe" 启动新版',
    '4. 新版启动瞬间弹出：',
]:
    bullet(t)
code_block("Error\nFailed to load Python DLL\n'C:\\Users\\28253\\AppData\\Local\\Temp\\_MEI331602\\python313.dll'\nLoadLibrary: 找不到指定的模块")
for t in [
    '5. _MEI 目录名每次失败都不同（_MEI72202、_MEI331602、_MEI246522...），点确定后目录被自动清理（Temp 里无残留）',
    '6. 必现：每次更新重启必失败（已复现 5 次以上）',
]:
    bullet(t)

# ============ 核心矛盾 ============
h1('三、最诡异的核心矛盾（重点求教）')
para('同一台电脑、同一个文件、同一套流程，用程序外的方式启动就成功，用户双击/更新重启就失败：')
rows = [
    ('测试', '结果'),
    ('用 subprocess 直接启动 v1.0.13 exe', '✅ 成功，窗口正常显示 v1.0.13'),
    ('用 os.startfile（等价资源管理器双击）启动 v1.0.13', '✅ 成功'),
    ('把 v1.0.13 复制到中文路径（F:\\论文排版工具_测试包\\）再启动', '✅ 成功'),
    ('完整模拟更新：v1.0.12 主程序运行中 → bat 强杀 → copy 替换 → start v1.0.13', '✅ 成功，新版正常启动'),
    ('清理 _MEI 残留后启动', '✅ 成功'),
    ('用户双击 v1.0.12', '✅ 成功（能正常打开）'),
    ('用户点自动更新 → bat 重启 v1.0.13', '❌ LoadLibrary 失败（必现）'),
    ('用户更新失败后手动双击 v1.0.13', '❌ LoadLibrary 失败'),
]
table = doc.add_table(rows=len(rows), cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b) in enumerate(rows):
    for j, val in enumerate((a, b)):
        cell = table.cell(i, j)
        cell.text = ''
        r = cell.paragraphs[0].add_run(val)
        r.font.name = '微软雅黑'
        r.font.size = Pt(9.5)
        if i == 0:
            r.bold = True
doc.add_paragraph()
para('即：同一台机器上，v1.0.12 双击能开，v1.0.13 双击/重启不能开；但本机（在用户机器上跑命令）启动 v1.0.13 全部成功。用户是交互式桌面会话正常双击。')

# ============ 已排查 ============
h1('四、已排查并排除的因素')
for t in [
    ('exe 文件损坏：', '排除。下载的 v1.0.13 md5 与发布源（GitHub release）完全一致；本机跑同一文件成功。'),
    ('杀软/Defender 拦截：', '排除。用户无杀软、防火墙关；多次启动成功无拦截；模拟"刚覆盖立即启动"也成功。'),
    ('中文路径：', '排除。英文路径（C:\\Temp）和中文路径（F:\\论文排版工具_测试包）均测试启动成功。'),
    ('_MEI 残留目录冲突：', '排除/已清理。曾发现 Temp 有 15 个 _MEI 残留（多次强杀遗留），已全部清理，并在 bat 中 start 前加清理、主程序启动时清 2 小时前残留。清理后仍失败。'),
    ('update.bat 逻辑缺陷：', '已修复。原 bat 是 taskkill 后立即 copy（无等待、无错误检查）；现改为 taskkill → 等 2s → copy（检查 errorlevel）→ 文件大小校验（<10MB 视为损坏）→ _MEI 清理 → 等 3s → start。修复后完整模拟（含主进程运行中）测试 bat 全链路：替换成功、新进程正常启动。用户实测仍失败。'),
    ('版本间打包差异：', 'v1.0.12 与 v1.0.13 用同一 spec、同一 PyInstaller 6.21.0、同一 conda 环境打包，只有 Python 源码差异（标准库级）。v1.0.12 用户能开、v1.0.13 不能开，但两个文件本机都能启动。'),
    ('Windows 事件日志：', 'Application 日志中无 DocFormatTool/python313/_MEI 相关错误记录（LoadLibrary 失败发生在 bootloader 早期 C 层，未触发 WER）。'),
    ('磁盘空间/权限：', 'C 盘剩 54.9GB、F 盘剩 221.9GB，Temp 目录用户可写（管理员）。'),
]:
    bullet(t[1], bold_prefix=t[0])

# ============ 技术特征 ============
h1('五、报错的技术特征')
for t in [
    "Failed to load Python DLL '<Temp>\\_MEIxxxxxx\\python313.dll' + LoadLibrary: 找不到指定的模块",
    '"找不到指定的模块"（ERROR_MOD_NOT_FOUND）通常指 python313.dll 自身加载时其依赖 DLL 缺失（不是 python313.dll 不存在）',
    '失败发生在 PyInstaller onefile bootloader 解压 _MEI 后、加载 python313.dll 时',
    '每次失败 _MEI 目录名不同（PID 或随机），说明每次都在正常解压新目录',
    '失败后点确定，_MEI 目录被清理（无残留可取证）',
]:
    bullet(t)

# ============ 疑问 ============
h1('六、疑问（请大神指点）')
for t in [
    '1. "找不到指定的模块" 具体缺的是哪个依赖？python313.dll 依赖 vcruntime140.dll、vcruntime140_1.dll（系统里都有），还有什么可能缺失？PyInstaller onefile 解压出的 python313.dll 依赖链里是否可能缺 ucrtbase 或 conda 特有 DLL？',
    '2. 为什么同一文件本机 subprocess/os.startfile 启动成功、用户双击/更新重启失败？两者在 Windows 进程创建、工作目录、环境变量、ShellExecute 参数上有什么本质差异？',
    '3. 为什么 v1.0.12（同一打包方式）用户双击能开、v1.0.13 不能开？打包内容差异（仅标准库源码差异）会导致 python313.dll 加载失败吗？',
    '4. PyInstaller 6.21.0 + Python 3.13 + conda 环境 + upx 的组合是否有已知的 LoadLibrary 问题？UPX 压缩是否可能导致解压后的 DLL 损坏/依赖异常？',
    '5. 有没有方法在 bootloader 阶段抓取真实的 LoadLibrary 错误码/缺失依赖（如启用 PyInstaller 的 bootloader 调试输出、或进程监控工具如 Process Monitor 看 _MEI 解压过程）？',
    '6. 是否建议直接放弃 onefile 换 onedir（DLL 放 exe 旁，不经 _MEI 解压）？onedir 在分发（用户下载 zip 解压）和自动更新（替换整个目录）上有什么成熟实践？',
]:
    bullet(t)

# ============ 补充发现 ============
h1('七、补充发现（排查过程中新线索）')
para('用户测试包目录残留 DocFormatTool_new.exe（2026/8/5 03:07），经验证 md5 与发布源 v1.0.13 完全一致——下载无损坏、copy 覆盖成功、重启启动的就是完整 v1.0.13。')
para('同时发现一个独立 bug：重写 update.bat 时遗漏了 copy 成功后删除 DocFormatTool_new.exe 的步骤，导致新版文件残留在安装目录（旧版 bat 有 del 逻辑）。该 bug 不影响启动，但需在后续版本修复。')

# ============ 下一步 ============
h1('八、下一步计划')
for t in [
    '1. 验证 onedir 打包（DLL 放 exe 旁，彻底绕开 _MEI 解压），若用户机器上 onedir 能稳定打开，则切换分发方式。',
    '2. 若 onedir 也失败，则用 Process Monitor 抓 bootloader 解压 + LoadLibrary 全过程定位缺失依赖。',
    '3. 修复 update.bat 遗漏的 del DocFormatTool_new.exe 步骤。',
]:
    bullet(t)

# 页脚说明
p = doc.add_paragraph()
r = p.add_run('生成时间：2026-08-05  |  程序：规范文档一键排版工具 DocFormatTool')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

out = r'LoadLibrary问题详细描述_求论坛.docx'
doc.save(out)
print('已生成:', out)
