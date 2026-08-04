# -*- coding: utf-8 -*-
"""生成 6 份成对样例（排版前/排版后）到新文件夹，供用户肉眼检查效果。
排版前 = 模拟学生手写格式（乱缩进/字号不一/手工编号）
排版后 = 走完整 pipeline 的输出"""
import os, sys, shutil
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from docx import Document
from docx.shared import Pt, Cm
from engine.config import load_preset
import engine.build_docx as B

OUT = r'F:\论文排版工具_测试包\人工检查样例'
if os.path.exists(OUT):
    shutil.rmtree(OUT)
os.makedirs(OUT)


def p(doc, text, size=12, bold=False, indent_cm=None, italic=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if indent_cm:
        para.paragraph_format.first_line_indent = Cm(indent_cm)
    return para


# ============ 01 数字体系（最常见） ============
d = Document()
p(d, '第一章  绪论', 22, True)
p(d, '随着人工智能技术的快速发展，教育领域正在经历深刻的变革。近年来，智能教育平台、自适应学习系统等应用层出不穷，为个性化教学提供了新的可能。', 12, indent_cm=0.74)
p(d, '本章首先介绍研究的背景与意义，然后对国内外相关研究进行综述，最后说明论文的组织结构。', 12, indent_cm=0.74)
p(d, '1.1  研究背景', 18, True)
p(d, '人工智能在教育领域的应用可以追溯到二十世纪八十年代的计算机辅助教学系统。进入二十一世纪后，随着深度学习技术的突破，智能教育进入快速发展阶段。', 12, indent_cm=0.74)
p(d, '1.1.1  国内研究现状', 16, True)
p(d, '国内学者对人工智能教育应用的研究主要集中在智能教学系统、教育大数据分析与个性化学习路径推荐等方面。', 12, indent_cm=0.74)
p(d, '1.1.2  国外研究现状', 16, True)
p(d, '国外研究起步较早，欧美发达国家在自适应学习系统与智能辅导系统方面积累了较为丰富的实证研究经验。', 12, indent_cm=0.74)
p(d, '1.2  研究意义', 18, True)
p(d, '本研究旨在探索人工智能技术在教学实践中的有效应用路径，为教育信息化建设提供理论参考与实践指导。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '01_数字体系_排版前.docx'))

# ============ 02 中文公文体系 ============
d = Document()
p(d, '一、总体概述', 22, True)
p(d, '本文围绕高校思想政治教育的实践创新展开研究，旨在探索新时代背景下思政教育的新模式与新路径。', 12, indent_cm=0.74)
p(d, '（一）研究背景', 18, True)
p(d, '随着社会环境的深刻变化，高校思想政治教育工作面临新的挑战与机遇，亟需在内容与方法上实现创新突破。', 12, indent_cm=0.74)
p(d, '1. 政策背景', 16, True)
p(d, '近年来，国家相继出台多项政策文件，对高校思想政治教育提出明确要求，为相关研究提供了政策依据。', 12, indent_cm=0.74)
p(d, '2. 现实需求', 16, True)
p(d, '当代大学生的思想特点和成长需求发生显著变化，传统教育模式难以完全适应新形势的要求。', 12, indent_cm=0.74)
p(d, '（二）研究意义', 18, True)
p(d, '本研究对于丰富思想政治教育理论、提升教育实效性具有重要的理论价值与实践意义。', 12, indent_cm=0.74)
p(d, '二、核心概念界定', 22, True)
p(d, '本章对研究中涉及的核心概念进行界定，明确研究的边界与范围。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '02_中文公文_排版前.docx'))

# ============ 03 混合体系（第一章 + 一、 + （一）） ============
d = Document()
p(d, '第一章  绪论', 22, True)
p(d, '本章是论文的总体介绍部分，阐述研究的问题、思路与方法。', 12, indent_cm=0.74)
p(d, '一、研究背景', 20, True)
p(d, '当前人工智能技术在各行业的渗透日益深入，教育领域同样面临技术变革带来的机遇与挑战。', 12, indent_cm=0.74)
p(d, '（一）国内现状', 16, True)
p(d, '国内智能教育起步虽晚但发展迅速，涌现出一批具有代表性的智能教育平台与产品。', 12, indent_cm=0.74)
p(d, '（二）国外现状', 16, True)
p(d, '国外在智能教育领域的研究更加系统深入，形成了较为完善的理论体系与应用模式。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '03_混合体系_排版前.docx'))

# ============ 04 无编号标题 ============
d = Document()
p(d, '引言', 22, True)
p(d, '论文写作是高等教育的重要环节，规范的格式是论文质量的基本保障。本文介绍一套自动排版工具的设计与实现。', 12, indent_cm=0.74)
p(d, '一、需求分析', 20, True)
p(d, '学生毕业论文格式不规范的问题普遍存在，人工调整费时费力且容易出错。', 12, indent_cm=0.74)
p(d, '结论', 22, True)
p(d, '本文提出的自动排版方案能够有效解决论文格式规范化问题，具有较好的实用价值。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '04_无编号标题_排版前.docx'))

# ============ 05 Markdown 标记 ============
d = Document()
p(d, '# 一级标题', 20, True)
p(d, '这是使用 Markdown 标记的一级标题下方正文内容。', 12, indent_cm=0.74)
p(d, '## 二级标题', 18, True)
p(d, '这是二级标题下方的正文内容。', 12, indent_cm=0.74)
p(d, '### 三级标题', 16, True)
p(d, '这是三级标题下方的正文内容。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '05_MD标记_排版前.docx'))

# ============ 06 格式混乱（模拟真实学生文档：缩进乱/字号乱/颜色） ============
d = Document()
p(d, '一、引言', 20, True)
p(d, '    随着互联网的普及，在线教育迅速发展。', 12, indent_cm=1.0)
p(d, '（一）背景', 16, True)
p(d, '在线教育打破了时空限制，为学习者提供了更加灵活的学习方式。', 12, indent_cm=0.5)
p(d, '1. 具体问题', 14, True)
p(d, '然而在线教育也面临学习效果难以保障、互动性不足等问题。', 12, indent_cm=0.74)
p(d, '二、主体', 20, True)
p(d, '本章针对上述问题展开深入分析。', 12, indent_cm=0.74)
d.save(os.path.join(OUT, '06_格式混乱_排版前.docx'))

# ============ 全部走 pipeline 排版 ============
print('排版中...')
for i in range(1, 7):
    name = ['01_数字体系', '02_中文公文', '03_混合体系', '04_无编号标题', '05_MD标记', '06_格式混乱'][i-1]
    src = os.path.join(OUT, name + '_排版前.docx')
    dst = os.path.join(OUT, name + '_排版后.docx')
    B.reformat_existing(load_preset('bachelor_cn'), src, dst)
    print('  %s ✓' % name)

# 检查要点说明
with open(os.path.join(OUT, '检查要点.txt'), 'w', encoding='utf-8') as f:
    f.write('人工检查样例 — 每对"排版前/排版后"对比要点：\n\n')
    f.write('01 数字体系：第一章=标题1，1.1=标题2，1.1.1/1.1.2=标题3；正文首行缩进2字符\n')
    f.write('02 中文公文：一、=标题1，（一）=标题2，1./2.=标题3\n')
    f.write('03 混合体系：第一章/一、=标题1，（一）（二）=标题2\n')
    f.write('04 无编号标题：引言/结论 识别为标题1，一、=标题1\n')
    f.write('05 MD标记：#/##/### = 标题1/2/3\n')
    f.write('06 格式混乱：乱缩进/字号被统一，一、/（一）/1. 层级正确\n\n')
    f.write('通用检查：打开"样式集"看 标题1/2/3 是否分层；正文小四宋体两端对齐；页码页脚。\n')
print('完成 → ' + OUT)
