# -*- coding: utf-8 -*-
"""v1.0.16 小红书文案 docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
from docx.oxml.ns import qn
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h1(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return p

def h2(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x09, 0x84, 0xE3)
    return p

def body(t, color=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    if color:
        r.font.color.rgb = color
    return p

doc.add_heading('v1.0.16 小红书宣传文案', level=0)
body('（以下文案可直接复制到小红书发布，图片用同目录 4 张宣传图）', RGBColor(0x94, 0xA3, 0xB8))

h1('【标题候选】')
body('1. 论文排版工具 v1.0.16 大更新！界面终于好看了 🎉')
body('2. 谁懂啊！论文排版工具居然有吉祥物了 🧸')
body('3. 论文排版工具焕新上线：拖拽选文件 + 全新界面')

h1('【正文】')
body('家人们谁懂啊，论文排版工具这次更新到 v1.0.16 了，界面直接焕然一新！')
body('')
h2('✨ 新界面')
body('侧边栏导航 + 卡片化布局，莫兰迪配色，告别老气的灰框框，一看就是正经软件该有的样子')
h2('📂 拖拽选文件')
body('论文文件直接拖进窗口，路径自动填好，不用再点来点去选半天')
h2('🖱️ 交互升级')
body('高级选项展开后内容再多也能鼠标滚轮滚动；窗口可以最大化；一键排版按钮永远固定在底部')
h2('🧸 专属吉祥物「排排」')
body('侧边栏常驻的小可爱，陪你写论文，萌到想多写两页')
body('')
h1('【这工具到底是干啥的】')
body('把 txt / md / docx 任意格式的文档，一键排成规范的 Word 论文格式：')
body('✅ 标题层级自动识别（第一章 / 1.1 / 一、/ （一）都能认）')
body('✅ 摘要、关键词、参考文献自动处理')
body('✅ 支持上传学校模板，按本校要求排版')
body('✅ 完全本地运行，论文不离开你的电脑')
body('✅ 免费、免安装、开箱即用')
body('')
h1('【获取方式】')
body('老用户：打开软件自动更新到 v1.0.16')
body('新用户：小红书 / 蓝奏云下载，免安装开箱即用')
body('')
h1('【标签】')
body('#论文排版 #毕业论文 #论文格式 #word排版 #大学生 #学术工具 #论文 #排版')

doc.save(r'site\xhs_out\小红书文案_v1016.docx')
print('文案 docx 已生成')
