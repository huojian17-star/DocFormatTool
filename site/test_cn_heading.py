# -*- coding: utf-8 -*-
"""测试中文公文编号体系层级映射：一、→H1，（一）→H2，1.→H3"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os, sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

# 构造测试文档：中文公文编号体系
d = Document()
paras = [
    ("一、总体概述", True),
    ("这是第一段正文内容，介绍总体情况。", False),
    ("（一）背景与意义", True),
    ("这是背景正文内容，讲研究背景和意义。", False),
    ("1. 研究背景", True),
    ("研究背景正文内容。", False),
    ("2. 研究意义", True),
    ("研究意义正文内容。", False),
    ("（二）国内外现状", True),
    ("国内外现状正文。", False),
    ("二、方法与过程", True),
    ("方法正文内容。", False),
    ("（一）数据来源", True),
    ("数据来源正文。", False),
    ("1. 问卷调查", True),
    ("问卷调查正文。", False),
    ("2. 访谈", True),
    ("访谈正文。", False),
]
for text, bold in paras:
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = __import__('docx').shared.Pt(22 if bold else 12)
    r.font.bold = bold

src = r'F:\论文排版工具_测试包\_cn_heading_test.docx'
d.save(src)

# 跑 pipeline（用通用模板）
from engine.config import load_preset
import engine.build_docx as B
cfg = load_preset('bachelor_cn')
out = r'F:\论文排版工具_测试包\_cn_heading_out.docx'
B.reformat_existing(cfg, src, out)

# 检查输出标题层级
w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
od = Document(out)
print('=== 中文编号层级映射结果 ===')
for p in od.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    pPr = p._p.find(w + 'pPr')
    ps = pPr.find(w + 'pStyle') if pPr is not None else None
    style = ps.get(w + 'val') if ps is not None else '(正文)'
    if style.startswith('Heading'):
        lvl = style[-1]
        print('  %-14s → Heading%s' % (t[:14], lvl))
