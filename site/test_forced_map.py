# -*- coding: utf-8 -*-
"""测试 forced_map：用与引擎一致的遍历（body.iter w:p + para_idx）建立索引"""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from engine.config import load_preset
import engine.build_docx as B

OUT = r'F:\论文排版工具_测试包\测试套件'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

src = os.path.join(OUT, 'T9_低置信_排版前.docx')
# 重建输入
d = Document()
def p(text, size=12):
    para = d.add_paragraph()
    r = para.add_run(text)
    r.font.size = Pt(size)
    return para
p('第一章 绪论', 22)
p('这是正文内容，用于测试低置信段落覆盖功能是否正常工作。', 12)
p('1. 研究背景', 14)
p('研究背景正文内容。', 12)
p('2. 研究意义', 14)
p('研究意义正文。', 12)
d.save(src)

# 与引擎一致：body.iter(w:p) 建立索引 → 文本映射
d2 = Document(src)
idx_map = {}
idx = 0
for p_el in d2.element.body.iter(qn('w:p')):
    t = ''.join(x.text or '' for x in p_el.iter(qn('w:t'))).strip()
    if t:
        idx_map[idx] = t
    idx += 1
print('索引映射:', {k: v[:10] for k, v in idx_map.items()})

# 找 "1. 研究背景" 和 "2. 研究意义" 的真实索引
i1 = next(k for k, v in idx_map.items() if v.startswith('1. 研究背景'))
i2 = next(k for k, v in idx_map.items() if v.startswith('2. 研究意义'))
forced = {"1. 研究背景": "body", "2. 研究意义": "heading2"}
print('设置 forced:', forced)

dst = os.path.join(OUT, 'T9_低置信_排版后.docx')
B.reformat_existing(load_preset('bachelor_cn'), src, dst, forced)

od = Document(dst)
print('=== forced_map 结果 ===')
for p_el in od.paragraphs:
    t = p_el.text.strip()
    if not t:
        continue
    pPr = p_el._p.find(W + 'pPr')
    ps = pPr.find(W + 'pStyle') if pPr is not None else None
    style = ps.get(W + 'val') if ps is not None else '(正文)'
    print('%-16s → %s' % (t[:16], style))
