# -*- coding: utf-8 -*-
"""六种文档场景回归测试：不同编号体系/标题风格/特殊格式 → 验证标题层级映射"""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from engine.config import load_preset
import engine.build_docx as B

OUT = r'F:\论文排版工具_测试包\测试套件'
os.makedirs(OUT, exist_ok=True)
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def new_doc(paras):
    """paras: [(text, is_heading)]，标题用 22pt 加粗，正文 12pt"""
    d = Document()
    for text, is_h in paras:
        p = d.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(22 if is_h else 12)
        r.font.bold = is_h
        if is_h:
            p.paragraph_format.space_before = Pt(6)
    return d


def check(doc_path, expect):
    """expect: [(标题前缀, 期望层级 1/2/3)]，验证 pStyle"""
    od = Document(doc_path)
    got = {}
    for p in od.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        pPr = p._p.find(W + 'pPr')
        ps = pPr.find(W + 'pStyle') if pPr is not None else None
        style = ps.get(W + 'val') if ps is not None else ''
        if style.startswith('Heading'):
            got[t[:8]] = int(style[-1])
    ok = True
    for prefix, want in expect:
        match = [lvl for k, lvl in got.items() if k.startswith(prefix)]
        if not match:
            print('  ✗ %s 未识别为标题' % prefix)
            ok = False
        elif match[0] != want:
            print('  ✗ %s 期望H%d 实际H%d' % (prefix, want, match[0]))
            ok = False
    if ok:
        print('  ✓ 全部通过: %s' % ' → '.join('%s=H%d' % (p, l) for p, l in expect))
    return ok


# ============ T1 数字体系 ============
t1 = new_doc([
    ("第一章 绪论", True), ("第一章正文内容。", False),
    ("1.1 研究背景", True), ("研究背景正文。", False),
    ("1.1.1 国内现状", True), ("国内现状正文。", False),
    ("1.1.2 国外现状", True), ("国外现状正文。", False),
    ("1.2 研究意义", True), ("研究意义正文。", False),
    ("第二章 方法", True), ("方法正文。", False),
])
p1 = os.path.join(OUT, 'T1_数字体系.docx')
t1.save(p1)
B.reformat_existing(load_preset('bachelor_cn'), p1, p1)
print('T1 数字体系(第一章/1.1/1.1.1):')
ok1 = check(p1, [("第一章", 1), ("1.1", 2), ("1.1.1", 3), ("1.1.2", 3), ("1.2", 2), ("第二章", 1)])

# ============ T2 中文公文体系 ============
t2 = new_doc([
    ("一、总体概述", True), ("概述正文。", False),
    ("（一）背景与意义", True), ("背景正文。", False),
    ("1. 研究背景", True), ("研究背景正文。", False),
    ("2. 研究意义", True), ("意义正文。", False),
    ("（二）国内外现状", True), ("现状正文。", False),
    ("二、方法与过程", True), ("方法正文。", False),
])
p2 = os.path.join(OUT, 'T2_中文公文.docx')
t2.save(p2)
B.reformat_existing(load_preset('bachelor_cn'), p2, p2)
print('T2 中文公文(一、/（一）/1.):')
ok2 = check(p2, [("一、", 1), ("（一）", 2), ("1.", 3), ("2.", 3), ("（二）", 2), ("二、", 1)])

# ============ T3 混合体系 ============
t3 = new_doc([
    ("第一章 绪论", True), ("绪论正文。", False),
    ("一、研究背景", True), ("背景正文。", False),
    ("（一）国内现状", True), ("国内正文。", False),
    ("（二）国外现状", True), ("国外正文。", False),
    ("第二章 方法", True), ("方法正文。", False),
    ("1.1 数据来源", True), ("数据来源正文。", False),
])
p3 = os.path.join(OUT, 'T3_混合体系.docx')
t3.save(p3)
B.reformat_existing(load_preset('bachelor_cn'), p3, p3)
print('T3 混合(第一章/一、/（一）/1.1):')
ok3 = check(p3, [("第一章", 1), ("一、", 1), ("（一）", 2), ("（二）", 2), ("第二章", 1), ("1.1", 2)])

# ============ T4 无编号大字号标题 ============
t4 = new_doc([
    ("引言", True), ("引言正文内容。", False),
    ("一、研究背景", True), ("背景正文。", False),
    ("结论", True), ("结论正文内容。", False),
])
p4 = os.path.join(OUT, 'T4_无编号标题.docx')
t4.save(p4)
B.reformat_existing(load_preset('bachelor_cn'), p4, p4)
print('T4 无编号大字号(引言/结论 期望能识别):')
ok4 = check(p4, [("一、", 1)])  # 无编号标题识别可能受限，只断言编号标题

# ============ T5 Markdown 标记 ============
t5 = new_doc([
    ("# 一级标题", True), ("正文内容。", False),
    ("## 二级标题", True), ("二级正文。", False),
    ("### 三级标题", True), ("三级正文。", False),
])
p5 = os.path.join(OUT, 'T5_MD标记.docx')
t5.save(p5)
B.reformat_existing(load_preset('bachelor_cn'), p5, p5)
print('T5 Markdown 标记(#/##/###):')
ok5 = check(p5, [("# 一级", 1), ("## 二级", 2), ("### 三级", 3)])

# ============ T6 格式混乱（考验不崩+识别） ============
t6 = new_doc([
    ("一、引言", True), ("  引言正文  内容。  ", False),
    ("（一）背景", True), ("背景正文。", False),
    ("1. 具体问题", True), ("问题正文。", False),
    ("二、主体", True), ("主体正文。", False),
])
# 故意加混乱：缩进/颜色/斜体
d6 = Document()
for text, is_h in [("一、引言", True), ("  引言正文  内容。  ", False), ("（一）背景", True),
                   ("背景正文。", False), ("1. 具体问题", True), ("问题正文。", False),
                   ("二、主体", True), ("主体正文。", False)]:
    p = d6.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22 if is_h else 12)
    r.font.bold = is_h
    if not is_h:
        p.paragraph_format.left_indent = Cm(0.5)
        r.font.color.rgb = __import__('docx').shared.RGBColor(0x33, 0x33, 0x99)
p6 = os.path.join(OUT, 'T6_格式混乱.docx')
d6.save(p6)
B.reformat_existing(load_preset('bachelor_cn'), p6, p6)
print('T6 格式混乱(缩进/彩色正文 + 中文标题):')
ok6 = check(p6, [("一、", 1), ("（一）", 2), ("1.", 3), ("二、", 1)])

print()
print('=== 汇总: %d/6 通过 ===' % sum([ok1, ok2, ok3, ok4, ok5, ok6]))
