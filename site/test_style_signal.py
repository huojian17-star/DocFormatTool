# -*- coding: utf-8 -*-
"""测试：输入文档带 Heading 样式/大纲级别时，强信号直接采用（不再靠正则猜）"""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from docx import Document
from docx.shared import Pt
from engine.config import load_preset
import engine.build_docx as B

OUT = r'F:\论文排版工具_测试包\测试套件'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# 构造：正文段落故意"看起来像标题"（短行+数字编号），但用 Normal；另一些用正确 Heading 样式
d = Document()
def p(text, style=None, size=12):
    para = d.add_paragraph(style=style) if style else d.add_paragraph()
    r = para.add_run(text)
    r.font.size = Pt(size)
    return para

# 正常标题（带样式）→ 强信号应识别
p('第一章 绪论', 'Heading 1', 22)
p('这是绪论正文内容，用于测试。', None, 12)
p('1.1 研究背景', 'Heading 2', 18)
p('研究背景正文。', None, 12)
# 陷阱段落：无样式、短行、数字编号——本应误判为标题，但前面有真实结构
p('1. 这里的数据来源于国家统计局：https://data.stats.gov.cn/', None, 12)
p('2. 资料来源：教育部历年报告。', None, 12)
# 大纲级别直接标记（无样式，手工 outlineLvl）
p3 = d.add_paragraph()
r = p3.add_run('手工大纲级别标题')
r.font.size = Pt(16)
pPr = p3._p.get_or_add_pPr()
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
ol = OxmlElement('w:outlineLvl')
ol.set(qn('w:val'), '1')  # 大纲级别2 → 标题2
pPr.append(ol)

src = os.path.join(OUT, 'T8_强信号_排版前.docx')
dst = os.path.join(OUT, 'T8_强信号_排版后.docx')
d.save(src)
B.reformat_existing(load_preset('bachelor_cn'), src, dst)

od = Document(dst)
print('=== 强信号测试结果 ===')
for p_el in od.paragraphs:
    t = p_el.text.strip()
    if not t:
        continue
    pPr = p_el._p.find(W + 'pPr')
    ps = pPr.find(W + 'pStyle') if pPr is not None else None
    style = ps.get(W + 'val') if ps is not None else '(正文)'
    print('%-16s → %s' % (t[:16], style))
