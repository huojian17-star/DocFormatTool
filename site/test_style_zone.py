# -*- coding: utf-8 -*-
"""测试：摘要长句/摘要内容/关键词/参考文献样式 + 脚注不误判标题"""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from docx import Document
from docx.shared import Pt, Cm
from engine.config import load_preset
import engine.build_docx as B

OUT = r'F:\论文排版工具_测试包\测试套件'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

d = Document()
def p(text, size=12, bold=False, indent=None):
    para = d.add_paragraph()
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if indent:
        para.paragraph_format.first_line_indent = Cm(indent)
    return para

# 标题
p('人工智能技术在教育领域的应用研究', 22, True)
# 摘要长句（标题+内容同一段）
p('摘  要：随着人工智能技术的快速发展，教育领域正在经历深刻变革。智能教育平台与自适应学习系统为个性化教学提供了新的可能，本研究对此展开系统分析。', 12)
# 摘要内容段（独立段）
p('本研究采用文献分析与案例研究相结合的方法，对人工智能教育应用现状进行梳理。', 12)
# 关键词
p('关键词：人工智能；教育应用；个性化学习；智能教育平台', 12)
# 正文标题
p('一、引言', 20, True)
p('人工智能技术在教育领域的应用日益广泛，本章介绍研究背景与意义。', 12)
# 脚注（应识别为正文/注释，不是标题3）
p('1. 数据来源于国家统计局：https://data.stats.gov.cn/，访问日期2026年8月。', 12)
p('2. 资料来源：教育部历年教育信息化发展报告。', 12)
# 二级标题 + 三级
p('（一）研究背景', 16, True)
p('1. 政策环境', 14, True)
p('近年来国家出台了多项支持教育信息化的政策文件。', 12)
# 参考文献标题
p('参考文献', 20, True)
# 参考文献条目
p('[1] 张伟, 李明. 人工智能教育应用研究综述[J]. 电化教育研究, 2025, 46(3): 10-18.', 12)
p('[2] 王芳. 教育数字化转型的路径与策略[M]. 北京: 高等教育出版社, 2024: 56-78.', 12)

src = os.path.join(OUT, 'T7_样式与脚注_排版前.docx')
dst = os.path.join(OUT, 'T7_样式与脚注_排版后.docx')
d.save(src)
B.reformat_existing(load_preset('bachelor_cn'), src, dst)

od = Document(dst)
print('=== 样式检查 ===')
for p_el in od.paragraphs:
    t = p_el.text.strip()
    if not t:
        continue
    pPr = p_el._p.find(W + 'pPr')
    ps = pPr.find(W + 'pStyle') if pPr is not None else None
    style = ps.get(W + 'val') if ps is not None else '(正文)'
    print('%-14s 样式=%-10s' % (t[:14], style))

print()
print('=== 样式表（styles.xml）中自定义样式 ===')
for st in od.styles:
    if st.type is not None and 'PARAGRAPH' in str(st.type) and st.name in ('摘要', '关键词', '参考文献'):
        print('  自定义样式存在:', st.name)
