# -*- coding: utf-8 -*-
"""决定性验证：模拟用户"城乡融合"文档结构（标题/摘要/关键词），分别测 txt(build) 和 docx(reformat) 流程，
检查输出 docx 里 标题/摘要标签/摘要正文/关键词 的 pStyle"""
import sys, os, tempfile
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))  # 项目根
from engine import infer, build_docx, config
from docx import Document

CFG = config.load_preset('bachelor_cn')

CONTENT = """城乡融合发展中的要素流动与资源配置优化研究
张三 李四
摘要：城乡融合发展是新时代中国式现代化的重要路径，本文从要素流动与资源配置的角度展开分析。
关键词：城乡融合；要素流动；资源配置
一、引言
城乡融合发展涉及人口、土地、资本等多维要素的自由流动。
二、理论基础
（一）要素流动理论
要素流动是区域协调发展的核心机制。
三、结论
城乡融合的关键在于破除体制机制障碍。
"""

TMP = os.path.join(tempfile.gettempdir(), 'cxrh_test')
os.makedirs(TMP, exist_ok=True)

def check(label, out_path, expect_styles):
    doc = Document(out_path)
    print('==== %s ====' % label)
    for p in doc.paragraphs[:8]:
        t = p.text.strip()
        if t:
            print('  [%s] %s' % (p.style.name, t[:28]))

# 1) txt → build()
txt = os.path.join(TMP, 'cxrh.txt')
open(txt, 'w', encoding='utf-8').write(CONTENT)
out1 = os.path.join(TMP, 'cxrh_txt_out.docx')
structs = infer.parse_file(txt)
infer._mark_cover(structs)
build_docx.build(CFG, structs, out1)
check('txt → build()', out1, None)

# 2) docx → reformat_existing()
# 构造：标题/摘要/关键词都是"正文样式"的普通段落（模拟用户原始文档）
src_docx = os.path.join(TMP, 'cxrh_src.docx')
from docx import Document as D2
d = D2()
for line in CONTENT.split('\n'):
    if line.strip():
        d.add_paragraph(line.strip())
d.save(src_docx)
out2 = os.path.join(TMP, 'cxrh_docx_out.docx')
build_docx.reformat_existing(CFG, src_docx, out2)
check('docx → reformat_existing()', out2, None)
