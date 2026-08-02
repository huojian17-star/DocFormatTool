# -*- coding: utf-8 -*-
"""生成《效果对比_最新版》文件夹：原文 + 最新引擎排版后（含标题样式）+ 报告。"""
import os
import shutil
import sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from engine import build_docx, config
from engine.styles import _set_run_font

OUT = r"F:\论文排版工具_测试包\效果对比_最新版"
os.makedirs(OUT, exist_ok=True)
for f in os.listdir(OUT):
    os.remove(os.path.join(OUT, f))

JOBS = [
    (r"D:\论文_管理研究方法 管理2304邓恺恒2302010189\管理2304邓恺恒2302010189.docx",
     "01_中文论文_管理研究方法", "bachelor_cn", "7章17节7表格"),
    (r"samples\real_papers\CN_Gostyan_docx-skill-4-cn-paper__测试论文.docx",
     "02_中文论文_测试论文", "bachelor_cn", "41表格6图"),
    (r"samples\real_papers\CN_hang660_lab-report-expert-SKILLS__example2_嵌入式计算机系统实验四实验报告.docx",
     "03_中文报告_嵌入式实验", "bachelor_cn", "2.2MB，73图"),
]

for src, name, preset_id, note in JOBS:
    if not os.path.exists(src):
        continue
    cfg = config.load_preset(preset_id)
    src_copy = os.path.join(OUT, name + "_原文.docx")
    shutil.copy2(src, src_copy)
    out = os.path.join(OUT, name + "_已排版.docx")
    try:
        stats = build_docx.reformat_existing(cfg, src_copy, out)
        build_docx.build_change_report(stats, cfg, src_copy, out)
        # 统计标题样式进入情况
        d = Document(out)
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        h_styles = {}
        for p_el in d.element.body.iter(ns + "p"):
            pPr = p_el.find(ns + "pPr")
            if pPr is not None:
                ps = pPr.find(ns + "pStyle")
                if ps is not None:
                    v = ps.get(ns + "val")
                    if v and v.startswith("Heading"):
                        h_styles[v] = h_styles.get(v, 0) + 1
        print("%s: 标题样式 %s" % (name, h_styles))
    except Exception as e:
        print("%s: 失败 %s" % (name, type(e).__name__))

# 汇总说明
doc = Document()
CN_B, CN_H, EN = "宋体", "黑体", "Times New Roman"


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_font(r, CN_H, EN, 15, bold=True)


def body(text):
    p = doc.add_paragraph(text)
    if p.runs:
        _set_run_font(p.runs[0], CN_B, EN, 11)
    p.paragraph_format.line_spacing = 1.4


h1("效果对比说明（最新引擎 v1.0.0）")
body("本文件夹用最新引擎排版，重点验证两个修复：")
body("1. 标题已进入 Word 样式集：打开《已排版》文件，按 Ctrl+Alt+1/2/3 或看左侧样式面板，能看到【标题 1/2/3】；视图-导航窗格可按标题跳转。")
body("2. 识别增强：第N章（阿拉伯数字）、（一）括号标题、参考文献条目（含[J]/[M]标识）、英文摘要/关键词均能识别。")
body("")
body("每篇含：原文（排版前）/ 已排版（排版后）/ 改动报告。对比看：标题样式、字体统一、表格/图片保留。")

doc.save(os.path.join(OUT, "查看说明.docx"))
print("文件夹已生成:", OUT)
