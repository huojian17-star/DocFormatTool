# -*- coding: utf-8 -*-
"""生成《第二曲线需求备忘_最终版.docx》：创作者专属 Agent 环境（启动即照此执行）。"""
import os
import sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from engine.styles import _set_run_font

CN_B, CN_H, EN = "宋体", "黑体", "Times New Roman"
OUT = r"F:\消失游戏开发日志\第二曲线_创作者Agent环境_需求备忘.docx"

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_font(r, CN_H, EN, 16, bold=True)
    p.paragraph_format.space_after = Pt(10)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, CN_H, EN, 13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def body(text, bold=False):
    p = doc.add_paragraph(text)
    if p.runs:
        _set_run_font(p.runs[0], CN_B, EN, 11, bold=bold)
    p.paragraph_format.line_spacing = 1.5


def item(text):
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        _set_run_font(p.runs[0], CN_B, EN, 11)
    p.paragraph_format.line_spacing = 1.4


h1("第二曲线产品：创作者专属 Agent 环境")
body("需求备忘（PRD v2 最终版）——本文件是未来启动项目的依据，启动时直接按此执行。")

h2("一、一句话定义")
body("对标 Codex / Claude Code / Reasonix 的 Agent 环境，但预置【创作调研 SOP】，"
     "天生懂考据类视频创作者的意图——不是通用助手，是创作者的专属调研协作者。", bold=True)

h2("二、需求来源（创作者亲历的痛点链，已逐层验证）")
item("剪视频最烦的不是剪辑，是搜集资料和写文案")
item("AI 联网搜索会幻觉 → 要来源链接 → 链接要一个个点开核对 → 挂了原文也不敢信 → 需要证据链")
item("AI 上下文有限 → 几万字考据稿只能生成几千字 → 穷尽性焦虑（怕漏）")
item("纯关键词抓取没有思考链 → 链式线索发现会漏（文章 A 提到新线索，爬虫不会顺藤摸瓜）")
item("本地 AI 对大众创作者负担过重（没电脑/低配）→ 云端是唯一大众化路径")
item("创作者把钱浪费在诉说需求上 → 准确理解创作意图是重中之重")

h2("三、产品形态（对标 agent 平台）")
body("创作者专属 Agent 环境：", bold=True)
item("交互层：终端/桌面界面")
item("模型层：用户 BYO key（OpenAI / DeepSeek / 本地 Ollama 可选）——零服务器成本，成本归用户")
item("核心资产：创作者 Skills 包（客制化工作流，即护城河）")

h2("四、核心资产：创作调研 SOP（意图理解是重中之重）")
body("创作者一句话需求 → 意图澄清 → 调研计划 → 链式执行 + 证据链标注", bold=True)
item("意图澄清：3-5 个问题（主题/体裁/观众预期/已有认知/想挖的层面）——问对问题比搜索对答案省钱")
item("调研计划：考据类视频自动展开五条线——人物/时间线/事件/争议/资料来源")
item("链式线索发现：读完文章 → 提炼新线索 → 建议下一步搜什么 → 循环（这是通用爬虫做不到的）")
item("素材证据链：AI 每句话旁标原出处 → 点链接精确定位原文段落 → 当场可验证")
item("素材覆盖率核查：素材库 vs 稿子逐条对照（47 份素材，3 份未覆盖）——量化回答‘会不会漏’")
item("网页内部文字读取 + 图片截取：自动存档，复用现有 OCR/vision 工具链")

h2("五、商业模式")
item("卖‘客制化工作流’（Skills 包），不是卖 AI")
item("BYO key：用户自带模型 key，你零服务器成本")
item("卖点 = 省 token：意图理解准 → 少跑偏 → 少烧钱 → 省下的钱 > 你收的钱（省出来的价值）")
item("与论文工具同构：密钥制、本地/云端可选、报告输出、自动更新")

h2("六、可行性评估")
item("意图澄清/SOP：产品设计问题，可做")
item("证据链/锚点定位：可做（本地 HTML 存档 + 锚点）")
item("链式线索发现：依赖模型推理能力（需实测本地或云端模型表现）")
item("零服务器成本：BYO key 成立")

h2("七、为什么现在不做（启动条件）")
item("论文工具是第一曲线：应优先验证‘卖本地工具’闭环并收第一笔钱")
item("本产品依赖 AI 推理质量 → 需先实测链式推理可行性")
body("启动条件（两个都满足）：", bold=True)
item("1. 论文工具月收入覆盖约 2 个月生活费（资金与模式验证）")
item("2. 用《午夜动物》素材实测：本地/云端模型能否完成链式线索发现（技术验证）")

h2("八、启动步骤（到时候照做）")
item("1. 复盘《午夜动物》素材流程：搜集→核对→整理的每一步耗时与痛点（产品规格来源）")
item("2. 用真实素材实测链式推理：AI 读文章→提炼线索→建议搜索→喂新网页，验证循环")
item("3. 设计意图澄清 SOP：找 3-5 个考据 UP 主访谈‘你希望 AI 先问你什么’")
item("4. 最小原型：意图澄清 + 证据链 + 覆盖率，复用论文工具体系")
item("5. 内测 → 收集考据党反馈 → 迭代")

h2("九、与论文工具的关系")
body("同一套技术底座（Python + 密钥 + 报告 + 自动更新）的两个产品。"
     "论文工具 = 单一功能工具（验证卖软件）；本产品 = 平台（多个 Skills，验证卖工作流）。"
     "论文排版未来可成为本环境的一个 Skill。")

h2("十、行业水位观察（入场时机信号）")
body("不启动开发，但监测云端推理成本：当‘完成一次完整考据链的云端成本’显著低于"
     "‘创作者愿意为省事付的钱’时（一次素材分析降至几分钱量级），才是入场时机。"
     "成本下跌是确定性趋势，需求已被验证，等水位。", bold=True)

doc.save(OUT)
print("已生成 ->", OUT)
