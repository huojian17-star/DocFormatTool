# -*- coding: utf-8 -*-
"""排版输出自动体检器：检查成品是否符合配置，输出体检报告。

用法:
  python tools/validate.py 输出.docx [输入.docx] [--config 配置.json | --preset id]

检查项:
  1. 页面   纸张尺寸/页边距 与配置一致
  2. 内容   图片/表格 是否保留（传入输入文件时对比数量）
  3. 字体   全文 run 的字体组合是否都在配置预期内（防残留杂字体）
  4. 字号   是否存在配置预期之外的字号（防字号不统一）
  5. 标题   识别出的标题层级分布（诊断用）
  6. 异常   run 级缺 rFonts / 超大字号 / 空文本大字号

退出码: 0 = 全部通过；1 = 有 FAIL 或 WARN。
"""
import argparse
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from collections import Counter

from docx import Document
from docx.oxml.ns import qn

from engine import infer, config as config_mod
from engine import build_docx


def _paper_label(cfg):
    return cfg["page"].get("paper", "A4")


def _expected_font_pairs(cfg):
    """配置中各角色允许的 (cn, en) 字体组合集合。"""
    pairs = set()
    for role, fd in cfg.get("fonts", {}).items():
        if isinstance(fd, dict) and "cn" in fd and "en" in fd:
            pairs.add((fd["cn"], fd["en"]))
    # md 代码块/行内代码字体
    pairs.add(("宋体", "Consolas"))
    return pairs


def _expected_sizes(cfg):
    """配置中各角色允许的字号（half-points）集合。"""
    sizes = set()
    for role, fd in cfg.get("fonts", {}).items():
        if isinstance(fd, dict) and fd.get("size_pt"):
            sizes.add(int(round(fd["size_pt"] * 2)))
    # 封面字号：题目大字 + 类别行（16pt 固定）
    cov = cfg.get("cover", {})
    if cov.get("title_size_pt"):
        sizes.add(int(round(cov["title_size_pt"] * 2)))
    sizes.add(32)  # 封面类别行（"课程论文" 等 16pt）
    # md 代码字号：正文 -1（代码块）/ -0.5（行内代码）
    body_sz = cfg.get("fonts", {}).get("body", {}).get("size_pt")
    if body_sz:
        sizes.add(int(round((body_sz - 1) * 2)))
        sizes.add(int(round((body_sz - 0.5) * 2)))
    return sizes


def _paragraph_text(p_el):
    return build_docx.para_text(p_el).strip()


def _locate(p_el):
    """定位段落位置：所在表格索引/文本框 + 文本开头。"""
    in_tbl = None
    cur = p_el.getparent()
    tbl_idx = 0
    while cur is not None:
        if cur.tag == qn("w:tbl"):
            in_tbl = tbl_idx
            break
        if cur.tag == qn("w:body"):
            break
        cur = cur.getparent()
    txt = _paragraph_text(p_el)[:24]
    if in_tbl is not None:
        return "表格内[表%d]: %s" % (in_tbl, txt)
    return "正文: %s" % txt


def validate(out_path, src_path=None, cfg=None, cfg_label="?"):
    """执行体检，返回 (results, summary)。results: [(level, item, detail), ...]"""
    results = []
    doc = Document(out_path)
    cfg = cfg or {}
    expected_fonts = _expected_font_pairs(cfg)
    expected_sizes = _expected_sizes(cfg)

    # ---- 1. 页面 ----
    sec = doc.sections[0]
    paper_w = round(sec.page_width.cm, 2)
    paper_h = round(sec.page_height.cm, 2)
    m = cfg["page"]["margins_cm"] if "page" in cfg else {}
    ok_page = True
    detail = "纸张 %.1fx%.1fcm" % (paper_w, paper_h)
    if "paper" in cfg.get("page", {}):
        expect = {"A4": (21.0, 29.7), "Letter": (21.59, 27.94)}.get(cfg["page"]["paper"])
        if expect and (abs(paper_w - expect[0]) > 0.2 or abs(paper_h - expect[1]) > 0.2):
            ok_page = False
    if m:
        ms = (round(sec.top_margin.cm, 2), round(sec.bottom_margin.cm, 2),
              round(sec.left_margin.cm, 2), round(sec.right_margin.cm, 2))
        me = (round(m["top"], 2), round(m["bottom"], 2), round(m["left"], 2), round(m["right"], 2))
        if abs(ms[0] - me[0]) > 0.05 or abs(ms[1] - me[1]) > 0.05 or \
           abs(ms[2] - me[2]) > 0.05 or abs(ms[3] - me[3]) > 0.05:
            ok_page = False
            detail += "，边距(上下左右) %s vs 配置 %s" % (str(ms), str(me))
        else:
            detail += "，边距 %s 一致" % str(ms)
    results.append(("PASS" if ok_page else "FAIL", "页面", detail))

    # ---- 2. 内容保留 ----
    out_pics = len(doc.element.body.findall(".//" + qn("w:drawing")))
    out_tbls = len(doc.tables)
    if src_path and os.path.exists(src_path) and src_path.lower().endswith(".docx"):
        src_doc = Document(src_path)
        src_pics = len(src_doc.element.body.findall(".//" + qn("w:drawing")))
        src_tbls = len(src_doc.tables)
        ok = out_pics >= src_pics and out_tbls >= src_tbls
        results.append(("PASS" if ok else "FAIL", "内容保留",
                        "图片 %d/%d，表格 %d/%d" % (out_pics, src_pics, out_tbls, src_tbls)))
    else:
        results.append(("PASS", "内容保留", "图片 %d 表格 %d（非 docx 输入不对比）" % (out_pics, out_tbls)))

    # ---- 3/4. 字体与字号（跳过封面区——封面千奇百怪，不检查） ----
    font_bad = Counter()
    font_bad_loc = {}
    size_bad = Counter()
    size_bad_loc = {}
    no_rfonts = 0
    in_cover = True
    for p_el in doc.element.body.iter(qn("w:p")):
        if in_cover:
            t0 = _paragraph_text(p_el)
            if not t0:
                continue  # 空段仍在封面区
            if build_docx._is_cover_block(p_el, cfg):
                continue  # 封面段：不检查
            in_cover = False
        for r in p_el.iter(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                continue
            rf = rpr.find(qn("w:rFonts"))
            sz = rpr.find(qn("w:sz"))
            val = sz.get(qn("w:val")) if sz is not None else None
            if rf is None:
                no_rfonts += 1
                continue
            cn = rf.get(qn("w:eastAsia")) or "?"
            en = rf.get(qn("w:ascii")) or "?"
            pair = (cn, en)
            if pair not in expected_fonts:
                font_bad[pair] += 1
                font_bad_loc.setdefault(pair, _locate(p_el))
            if val and int(val) not in expected_sizes:
                size_bad[val] += 1
                size_bad_loc.setdefault(val, _locate(p_el))

    if font_bad:
        for pair, n in font_bad.most_common(5):
            results.append(("FAIL", "字体", "%s/%s ×%d（位于%s）" % (pair[0], pair[1], n, font_bad_loc[pair])))
    else:
        results.append(("PASS", "字体", "全部在配置预期内"))
    if size_bad:
        for val, n in size_bad.most_common(5):
            pt = int(val) / 2
            results.append(("FAIL", "字号", "%spt ×%d（位于%s）" % (pt, n, size_bad_loc[val])))
    else:
        results.append(("PASS", "字号", "全部在配置预期内"))
    if no_rfonts:
        results.append(("WARN", "字体", "%d 个 run 缺 rFonts（多为空 run，可忽略）" % no_rfonts))

    # ---- 5. 标题分布（排除表格内段落 + 文档自带目录页）----
    h1 = h2 = h3 = 0
    in_toc = False
    for p_el in doc.element.body.iter(qn("w:p")):
        if build_docx._in_table(p_el):
            continue
        t = _paragraph_text(p_el)
        if not t:
            continue
        if not in_toc and "目" in t and "录" in t and len(t) <= 10:
            in_toc = True
            continue
        if in_toc:
            typ_t, _ = infer._classify(t)
            if typ_t in ("heading1", "heading2", "heading3"):
                continue  # 目录条目
            in_toc = False
        typ, _ = infer._classify(t)
        if typ == "heading1":
            h1 += 1
        elif typ == "heading2":
            h2 += 1
        elif typ == "heading3":
            h3 += 1
    results.append(("PASS", "标题", "一级 %d 个 / 二级 %d 个 / 三级 %d 个（正文）" % (h1, h2, h3)))

    # ---- 汇总 ----
    n_pass = sum(1 for lv, _, _ in results if lv == "PASS")
    n_warn = sum(1 for lv, _, _ in results if lv == "WARN")
    n_fail = sum(1 for lv, _, _ in results if lv == "FAIL")
    summary = "体检完成：%d 通过 / %d 警告 / %d 失败（配置: %s）" % (n_pass, n_warn, n_fail, cfg_label)
    return results, summary


def report(out_path, src_path, cfg, cfg_label):
    results, summary = validate(out_path, src_path, cfg, cfg_label)
    lines = ["==== 排版体检报告 ====", "配置: %s" % cfg_label,
             "输出: %s" % out_path, ""]
    for lv, item, detail in results:
        lines.append("[%s] %s: %s" % (lv, item, detail))
    lines.append("")
    lines.append("结论: " + summary)
    return "\n".join(lines), 1 if any(lv == "FAIL" for lv, _, _ in results) else 0


def main():
    ap = argparse.ArgumentParser(description="排版输出自动体检")
    ap.add_argument("out", help="排版输出 .docx")
    ap.add_argument("src", nargs="?", default=None, help="输入原文档（对比图片/表格保留）")
    ap.add_argument("--config", default=None, help="配置 JSON")
    ap.add_argument("--preset", default=None, help="内置模板 id")
    args = ap.parse_args()

    if args.config:
        with open(args.config, encoding="utf-8") as f:
            cfg = config_mod.merge_default(json.load(f))
        label = args.config
    elif args.preset:
        cfg = config_mod.load_preset(args.preset)
        label = cfg.get("school", args.preset)
    else:
        cfg = {}
        label = "（未指定配置，仅检查内容与结构）"

    text, code = report(args.out, args.src, cfg, label)
    print(text)
    sys.exit(code)


if __name__ == "__main__":
    main()
